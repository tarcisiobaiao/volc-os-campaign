#!/usr/bin/env python3
"""Atualiza o Workbook VOLC O.S. para v0.3 com a camada Graphify."""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import gerar_workbook_volc_os as b
import gerar_workbook_volc_os_v02 as v02


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "workbook-volc-os" / "assets-v03"
OUT = ROOT / "entregaveis" / "Workbook_VOLC_OS_Mapa_Vivo_Graphify_e_Prioridades_v0.3.docx"
BASE = ROOT / "entregaveis" / "Workbook_VOLC_OS_Mapa_Mestre_e_Prioridades_v0.2.docx"
EXPLORER_SCREENSHOT = Path("/private/tmp/volc-neural-map-v03.png")


def save(image: Image.Image, name: str) -> Path:
    path = ASSET_DIR / name
    image.save(path, quality=96)
    return path


def cover() -> Path:
    image = Image.new("RGB", (1400, 1980), b.rgb(b.NAVY))
    light = Image.new("RGBA", image.size, (0, 0, 0, 0))
    draw_light = ImageDraw.Draw(light)
    for cx, cy, color, radius, alpha in [
        (1200, 360, b.CYAN, 380, 115),
        (1090, 650, b.PURPLE, 520, 105),
        (1240, 970, b.ORANGE, 390, 95),
        (70, 1790, b.DEEP, 470, 125),
    ]:
        draw_light.ellipse((cx-radius, cy-radius, cx+radius, cy+radius), fill=(*b.rgb(color), alpha))
    light = light.filter(ImageFilter.GaussianBlur(112))
    image = Image.alpha_composite(image.convert("RGBA"), light).convert("RGB")
    draw = ImageDraw.Draw(image)
    if b.LOGO_LIGHT.exists():
        logo = Image.open(b.LOGO_LIGHT).convert("RGBA")
        logo.thumbnail((390, 150))
        image.paste(logo, (90, 90), logo)
    draw.text((95, 455), "WORKBOOK · V0.3", font=b.font(38, display=True), fill=b.rgb(b.CYAN))
    draw.text((90, 540), "VOLC O.S.", font=b.font(112, display=True), fill=b.rgb(b.WHITE))
    draw.text((95, 690), "Mapa vivo", font=b.font(76, display=True), fill=b.rgb(b.WHITE))
    draw.text((95, 790), "+ próximos passos", font=b.font(70, display=True), fill=b.rgb(b.WHITE))
    b.text_wrap(draw, "Ver o todo. Encontrar o caminho. Executar sem se perder.",
                (100, 915), 1080, b.font(36), b.WHITE, 12)
    b.rounded(draw, (95, 1160, 1070, 1450), 28, "141D36", b.CYAN, 3)
    draw.text((135, 1205), "NOVA CAMADA DE LEITURA", font=b.font(24, bold=True), fill=b.rgb(b.CYAN))
    b.text_wrap(draw,
                "Mapa Mestre para decidir. Explorador Neural para navegar. Graphify para provar qual implementação sustenta cada relação.",
                (135, 1262), 845, b.font(28), b.WHITE, 10)
    draw.text((98, 1782), "Snapshot do sistema  •  22 de agosto de 2026", font=b.font(22), fill=b.rgb("B7C3D9"))
    draw.text((98, 1827), "9.161 nós  •  21.230 relações  •  98 pontes negócio → código", font=b.font(20), fill=b.rgb("B7C3D9"))
    return save(image, "00_capa_v03.png")


def three_layers() -> Path:
    image, draw = b.canvas(
        "Três níveis para não se perder",
        "Cada lente responde uma pergunta diferente. Nenhuma precisa carregar a complexidade das outras.",
        dark=True,
    )
    layers = [
        ("01", "DECIDIR", "Workbook", "O que vem primeiro, por quê e como provar a saída.", b.ORANGE),
        ("02", "NAVEGAR", "Mapa Mestre", "Capacidades, estados, dependências e evidências operacionais.", b.CYAN),
        ("03", "INVESTIGAR", "Explorador Neural", "Caminhos, vizinhança, implementação e impacto no código.", b.PURPLE),
    ]
    for index, (number, verb, title, desc, color) in enumerate(layers):
        y = 230 + index * 255
        draw.text((110, y + 20), number, font=b.font(62, display=True), fill=b.rgb(color))
        draw.line((260, y + 58, 1690, y + 58), fill=b.rgb("33405A"), width=2)
        draw.text((305, y + 8), verb, font=b.font(18, bold=True), fill=b.rgb(color))
        draw.text((305, y + 50), title, font=b.font(36, display=True), fill=b.rgb(b.WHITE))
        b.text_wrap(draw, desc, (760, y + 42), 850, b.font(24), "D5DDEA", 8)
    draw.text((110, 985), "A regra", font=b.font(18, bold=True), fill=b.rgb(b.ORANGE))
    draw.text((260, 974), "Comece amplo. Aproxime somente quando uma decisão pedir prova.",
              font=b.font(27, display=True), fill=b.rgb(b.WHITE))
    return save(image, "12_tres_niveis.png")


def trust_graph() -> Path:
    image, draw = b.canvas(
        "Hipótese nunca se disfarça de fato",
        "Toda relação do grafo profundo carrega procedência e nível de confiança.",
    )
    left = (105, 260, 830, 780)
    right = (970, 260, 1695, 780)
    b.rounded(draw, left, 34, b.PALE_BLUE, b.CYAN, 3)
    b.rounded(draw, right, 34, b.PALE_PURPLE, b.PURPLE, 3)
    draw.text((160, 315), "EXTRAÍDA", font=b.font(27, bold=True), fill=b.rgb(b.CYAN))
    draw.text((1025, 315), "INFERIDA", font=b.font(27, bold=True), fill=b.rgb(b.PURPLE))
    b.text_wrap(draw, "Apareceu diretamente em código, rota, banco ou inventário medido.",
                (160, 390), 590, b.font(29, display=True), b.NAVY, 10)
    b.text_wrap(draw, "Foi resolvida por correspondência ou modelada como relação de negócio.",
                (1025, 390), 590, b.font(29, display=True), b.NAVY, 10)
    for index, text in enumerate(["import", "chamada", "rota", "contagem", "arquivo"]):
        x = 160 + (index % 3) * 195
        y = 590 + (index // 3) * 68
        b.rounded(draw, (x, y, x + 168, y + 48), 15, b.WHITE, b.CYAN, 2)
        draw.text((x + 18, y + 13), text, font=b.font(17, bold=True), fill=b.rgb(b.DEEP))
    for index, text in enumerate(["alimenta", "habilita", "depende", "amadurece", "prioriza"]):
        x = 1025 + (index % 3) * 195
        y = 590 + (index // 3) * 68
        b.rounded(draw, (x, y, x + 168, y + 48), 15, b.WHITE, b.PURPLE, 2)
        draw.text((x + 18, y + 13), text, font=b.font(17, bold=True), fill=b.rgb(b.DEEP))
    b.arrow(draw, (835, 520), (955, 520), b.ORANGE, 8, 22)
    b.rounded(draw, (480, 865, 1320, 960), 24, b.NAVY)
    draw.text((535, 895), "Quando a evidência não basta: AMBÍGUA → revisão humana",
              font=b.font(24, display=True), fill=b.rgb(b.WHITE))
    return save(image, "13_confianca_grafo.png")


def replace_text(doc: Document) -> None:
    replacements = {
        "A versão 0.2 começa reconhecendo uma leitura incorreta da versão anterior.":
            "A versão 0.3 preserva a correção central e acrescenta uma camada navegável de prova.",
        "Este workbook v0.2": "Este workbook v0.3",
        "VOLC O.S.  •  Mapa Mestre e Prioridades  •  v0.2":
            "VOLC O.S.  •  Mapa vivo e Prioridades  •  v0.3",
        "VOLC O.S. — Mapa Mestre e Prioridades v0.2":
            "VOLC O.S. — Mapa vivo, Graphify e Prioridades v0.3",
        "Regra para a v0.3": "Regra para a próxima edição",
        "269 nós e 442 relações em um artefato navegável e legível por máquina.":
            "Mapa executivo: 269 nós e 442 relações. Grafo profundo: 9.161 nós, 21.230 relações e 98 pontes.",
    }
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def add_explorer_to_usage_table(doc: Document) -> None:
    for table in doc.tables:
        text = " ".join(cell.text for row in table.rows for cell in row.cells)
        if "Mapa_Mestre_VOLC_OS.html" not in text:
            continue
        cells = table.add_row().cells
        b.set_cell_shading(cells[0], "F7F9FC")
        b.set_cell_shading(cells[1], "F7F9FC")
        for cell in cells:
            b.set_cell_margins(cell, 105, 120, 105, 120)
            b.set_cell_border(cell, bottom={"val": "single", "sz": "5", "color": "DDE4EE"})
        first = cells[0].paragraphs[0].add_run("Explorador_Neural_VOLC_OS.html")
        first.font.size = Pt(8.5)
        first.font.color.rgb = RGBColor.from_string(b.INK)
        second = cells[1].paragraphs[0].add_run("Buscar, aproximar e atravessar o grafo híbrido de forma dinâmica.")
        second.font.size = Pt(8.5)
        second.font.color.rgb = RGBColor.from_string(b.INK)
        break


def remove_old_closing_page(doc: Document) -> None:
    body = doc._body._element
    tables = [element for element in body if element.tag == qn("w:tbl")]
    if tables:
        body.remove(tables[-1])
    for element in reversed(list(body)):
        if element.tag == qn("w:p") and element.xpath(".//w:br[@w:type='page']"):
            body.remove(element)
            break


def replace_cover(doc: Document, new_cover: Path) -> None:
    inline = doc.inline_shapes[0]._inline
    blip = inline.graphic.graphicData.pic.blipFill.blip
    image_part = doc.part.related_parts[blip.embed]
    image_part._blob = new_cover.read_bytes()


def add_graphify_section(doc: Document, images: dict[str, Path]) -> None:
    b.page_break(doc)
    b.add_section_divider(
        doc, "05", "Um sistema que pode ser atravessado",
        "O mapa deixou de ser apenas inventário. Agora ele permite aproximar, buscar, seguir relações e provar onde cada capacidade vive.",
        b.CYAN,
    )

    b.page_break(doc)
    b.add_title(doc, "O mapa agora tem profundidade",
                "A complexidade fica disponível quando necessária, sem ocupar a tela de decisão.",
                "Arquitetura de leitura")
    b.add_image(doc, images["layers"])
    b.add_caption(doc, "Figura 12 — Três níveis de leitura sobre a mesma fonte de verdade.")
    b.simple_table(doc, ["LENTE", "ABRA QUANDO", "SAIA COM"], [
        ("Workbook", "Você precisa ordenar ou aprovar trabalho.", "Uma decisão, pré-requisitos e prova de saída."),
        ("Mapa Mestre", "Você precisa entender o todo e o estado atual.", "Um recorte de capacidades e dependências."),
        ("Explorador Neural", "Você precisa investigar caminho, impacto ou implementação.", "Nós, relações e evidências concretas."),
    ], [4.1, 6.5, 6.3], font_size=8.6)

    b.page_break(doc)
    b.add_title(doc, "O Explorador Neural",
                "Uma experiência semelhante ao grafo do Obsidian, desenhada para o contexto operacional do VOLC O.S.",
                "Navegação dinâmica")
    if images.get("screenshot") and images["screenshot"].exists():
        b.add_image(doc, images["screenshot"])
        b.add_caption(doc, "Figura 13 — Lente Mapa: 269 nós operacionais. Zoom, busca, seleção e vizinhança funcionam localmente.")
    b.simple_table(doc, ["LENTE", "O QUE MOSTRA"], [
        ("1 · Mapa", "Somente a operação curada: capacidades, telas, banco, n8n, tarefas e documentos."),
        ("2 · Pontes", "A operação mais os arquivos e objetos técnicos que a implementam."),
        ("3 · Rede completa", "Os 9.161 nós e 21.230 relações do snapshot integral."),
        ("4 · Vizinhança", "Um nó selecionado e até dois saltos ao redor, para investigar sem ruído."),
    ], [4.2, 12.7])
    b.callout(doc, "Como abrir",
              "Dê duplo clique em entregaveis/Explorador_Neural_VOLC_OS.html. O arquivo é standalone: não precisa de servidor, login ou internet.",
              b.CYAN, b.PALE_BLUE)

    b.page_break(doc)
    b.add_title(doc, "Como explorar sem se perder",
                "Comece pela pergunta, não pela rede completa.",
                "Ritual de navegação")
    steps = [
        ("1", "BUSQUE", "Digite uma capacidade, tabela, workflow, tela ou arquivo."),
        ("2", "SELECIONE", "Leia o que é, estado, origem, evidência e grau de conectividade."),
        ("3", "APROXIME", "Use Focar 1 salto ou Focar 2 saltos para revelar apenas o necessário."),
        ("4", "COMPARE", "Confirme se a relação é extraída ou inferida antes de decidir."),
        ("5", "VOLTE", "Pressione Esc para regressar ao Mapa e recuperar o contexto global."),
    ]
    for number, title, description in steps:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        table.columns[0].width = Cm(1.1)
        table.columns[1].width = Cm(3.2)
        table.columns[2].width = Cm(12.6)
        b.set_cell_shading(table.cell(0, 0), b.NAVY)
        b.set_cell_shading(table.cell(0, 1), b.WHITE)
        b.set_cell_shading(table.cell(0, 2), b.WHITE)
        for cell in table.rows[0].cells:
            b.set_cell_margins(cell, 135, 145, 135, 145)
            b.set_cell_border(cell, bottom={"val": "single", "sz": "7", "color": "DDE4EE"})
        p = table.cell(0, 0).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(number)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(b.WHITE)
        p = table.cell(0, 1).paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Space Grotesk"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(b.CYAN)
        p = table.cell(0, 2).paragraphs[0]
        run = p.add_run(description)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(b.MUTED)
    b.callout(doc, "Atalhos", "/ busca  •  1–4 alternam a lente  •  F enquadra  •  Esc volta ao Mapa",
              b.DEEP, b.PALE_BLUE)

    b.page_break(doc)
    b.add_title(doc, "A trilha de confiança",
                "O Graphify trouxe uma regra simples que protege o roadmap de conclusões apressadas.",
                "Fato e hipótese")
    b.add_image(doc, images["trust"])
    b.add_caption(doc, "Figura 14 — Confiança explícita em cada relação.")
    b.simple_table(doc, ["MARCA", "LEITURA", "DECISÃO"], [
        ("EXTRACTED", "A relação apareceu diretamente na fonte.", "Pode sustentar diagnóstico, desde que a fonte esteja atual."),
        ("INFERRED", "A relação foi resolvida ou modelada.", "É uma hipótese útil; confirme antes de mudar arquitetura ou gasto."),
        ("AMBIGUOUS", "A evidência não diferencia as interpretações.", "Exige revisão humana e não deve virar automação."),
    ], [3.2, 6.5, 7.2], font_size=8.4)

    b.page_break(doc)
    b.add_title(doc, "O que esta camada mudou no roadmap",
                "Ela não inventou uma nova prioridade. Ela tornou as prioridades atuais demonstráveis.",
                "Consequência prática")
    b.simple_table(doc, ["ANTES", "AGORA"], [
        ("Documentos descreviam o sistema por recortes.", "O grafo encontra correspondências entre operação e implementação."),
        ("Uma prioridade podia parecer ausente porque a tela não era lembrada.", "Busca e caminhos mostram o que já existe antes de propor outra construção."),
        ("Relações modeladas pareciam tão certas quanto chamadas reais.", "Fato extraído e inferência recebem marcas diferentes."),
        ("Investigar impacto exigia reler muitos arquivos.", "Vizinhança e conexões reduzem a investigação ao recorte relevante."),
    ], [8.3, 8.6])
    b.callout(doc, "A prioridade continua",
              "Terminar a ponte Hub de Tráfego → Nascimento Search → Cockpit existente. O grafo agora permite provar esse caminho e localizar seus pontos de implementação.",
              b.ORANGE, b.PALE_ORANGE)


def add_closing(doc: Document) -> None:
    b.page_break(doc)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17.3)
    cell = table.cell(0, 0)
    b.set_cell_shading(cell, b.NAVY)
    b.set_cell_margins(cell, 850, 650, 850, 650)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Veja o todo.\n")
    run.font.name = "Space Grotesk"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(b.WHITE)
    run = paragraph.add_run("Aproxime para provar.\n\n")
    run.font.name = "Space Grotesk"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(b.CYAN)
    run = paragraph.add_run("VOLC O.S.  •  Mapa vivo, Graphify e Prioridades  •  v0.3")
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("B7C3D9")


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if not BASE.exists():
        v02.main()
    doc = Document(BASE)
    replace_text(doc)
    add_explorer_to_usage_table(doc)
    remove_old_closing_page(doc)
    images = {
        "cover": cover(),
        "layers": three_layers(),
        "trust": trust_graph(),
        "screenshot": EXPLORER_SCREENSHOT,
    }
    replace_cover(doc, images["cover"])
    add_graphify_section(doc, images)
    add_closing(doc)
    doc.core_properties.title = "VOLC O.S. — Mapa vivo, Graphify e Prioridades v0.3"
    doc.core_properties.subject = "Workbook executivo, grafo híbrido e roadmap operacional"
    doc.core_properties.author = "VOLC"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
