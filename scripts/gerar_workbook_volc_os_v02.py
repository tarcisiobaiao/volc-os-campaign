#!/usr/bin/env python3
"""Workbook VOLC O.S. v0.2 — reconstruído a partir do Mapa Mestre."""

from __future__ import annotations

import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

import gerar_workbook_volc_os as b


ROOT = Path(__file__).resolve().parents[1]
GRAPH_PATH = ROOT / "docs" / "volc-os-graph" / "volc-os-graph.json"
ASSET_DIR = ROOT / "docs" / "workbook-volc-os" / "assets-v02"
OUT = ROOT / "entregaveis" / "Workbook_VOLC_OS_Mapa_Mestre_e_Prioridades_v0.2.docx"
W, H = b.W, b.H


def save(im, name):
    path = ASSET_DIR / name
    im.save(path, quality=95)
    return path


def cover():
    im = Image.new("RGB", (1400, 1980), b.rgb(b.NAVY))
    layer = Image.new("RGBA", im.size, (0, 0, 0, 0)); ld = ImageDraw.Draw(layer)
    for cx, cy, color, radius in [(1180,420,b.CYAN,420),(1070,720,b.PURPLE,540),(1280,1030,b.ORANGE,400),(120,1750,b.DEEP,480)]:
        ld.ellipse((cx-radius,cy-radius,cx+radius,cy+radius),fill=(*b.rgb(color),110))
    layer=layer.filter(ImageFilter.GaussianBlur(115)); im=Image.alpha_composite(im.convert("RGBA"),layer).convert("RGB")
    d=ImageDraw.Draw(im)
    if b.LOGO_LIGHT.exists():
        logo=Image.open(b.LOGO_LIGHT).convert("RGBA"); logo.thumbnail((390,150)); im.paste(logo,(90,90),logo)
    d.text((95,470),"WORKBOOK · V0.2",font=b.font(38,display=True),fill=b.rgb(b.CYAN))
    d.text((90,545),"VOLC O.S.",font=b.font(112,display=True),fill=b.rgb(b.WHITE))
    d.text((95,690),"Mapa mestre",font=b.font(76,display=True),fill=b.rgb(b.WHITE))
    d.text((95,790),"e prioridades",font=b.font(76,display=True),fill=b.rgb(b.WHITE))
    b.text_wrap(d,"O sistema inteiro antes do próximo passo.",(100,920),1080,b.font(39),b.WHITE,12)
    b.rounded(d,(95,1170,1040,1435),28,"141D36",b.CYAN,3)
    d.text((135,1215),"CORREÇÃO DE LEITURA",font=b.font(25,bold=True),fill=b.rgb(b.CYAN))
    b.text_wrap(d,"O cockpit de campanha já existe. A prioridade é conectar, reconciliar e fechar os elos — não construir outro monitoramento.",(135,1270),810,b.font(29),b.WHITE,10)
    d.text((98,1785),"Snapshot do sistema  •  22 de agosto de 2026",font=b.font(22),fill=b.rgb("B7C3D9"))
    d.text((98,1830),"Fonte: código + Supabase vivo + n8n + ClickUp + documentos oficiais",font=b.font(20),fill=b.rgb("B7C3D9"))
    return save(im,"00_capa_v02.png")


def graph_overview(graph):
    im,d=b.canvas("O VOLC O.S. inteiro em uma página","24 capacidades de negócio organizadas em 10 clusters — o inventário total tem 269 nós.",dark=True)
    caps=[n for n in graph["nodes"] if n["type"]=="capability"]
    cluster_order=list(graph["clusters"])
    centers={c:(210+(i%5)*345,270+(i//5)*430) for i,c in enumerate(cluster_order)}
    state_color={"live":b.GREEN,"implemented":b.CYAN,"partial":b.ORANGE,"historical":b.PURPLE,"inactive":"667085","empty":"667085","decision":b.ORANGE,"risk":b.ORANGE,"todo":b.CYAN}
    for c,(cx,cy) in centers.items():
        d.text((cx-145,cy-135),graph["clusters"][c].upper(),font=b.font(15,bold=True),fill=b.rgb("7F91AD"))
        arr=[n for n in caps if n["cluster"]==c]
        for j,n in enumerate(arr):
            x=cx-145+(j%2)*150; y=cy-95+(j//2)*78; col=state_color.get(n["state"],"667085")
            b.rounded(d,(x,y,x+140,y+61),16,"151E35",col,2)
            label=n["label"] if len(n["label"])<22 else n["label"][:21]+"…"
            b.text_wrap(d,label,(x+10,y+13),120,b.font(15,bold=True),b.WHITE,3,max_lines=2)
    d.text((80,982),"Verde: vivo agora   Azul: implementado   Laranja: parcial/decisão/risco   Roxo: histórico   Cinza: vazio/inativo",font=b.font(19),fill=b.rgb("AAB5C5"))
    return save(im,"01_grafo_geral.png")


def product_surfaces():
    im,d=b.canvas("O produto que já existe","As telas abaixo não são promessa: estão registradas no aplicativo atual.")
    groups=[
        ("OPERAR",[("Dashboard geral","Economia do portfólio"),("Projetos","Visão por ativo"),("Campanhas","Lista e configuração"),("Campanha","Cockpit completo"),("Relatórios","Análises avançadas")],b.CYAN),
        ("PRODUZIR",[("Pautador Pro","Descoberta e validação"),("Redator","Funis e páginas"),("Publicação","Prova e WordPress"),("Incubadora","Fábrica de publishers")],b.PURPLE),
        ("ADQUIRIR",[("Hub de Tráfego","Fila e campanhas no ar"),("Nova Search","Copy, lance, prova e recibo"),("Alertas","Sino global e entrega")],b.ORANGE),
        ("GOVERNAR",[("Custos","Economia operacional"),("Integrações","Google e Meta"),("Usuários","Acesso e comissão"),("Admin v6","Novo RBAC em código")],b.DEEP),
    ]
    for gi,(tag,items,color) in enumerate(groups):
        x=85+gi*425
        d.text((x,225),tag,font=b.font(21,bold=True),fill=b.rgb(color))
        for j,(title,desc) in enumerate(items):
            y=275+j*135
            b.rounded(d,(x,y,x+385,y+108),22,b.WHITE,"D8DEE9",2)
            d.rectangle((x,y,x+9,y+108),fill=b.rgb(color))
            d.text((x+27,y+20),title,font=b.font(24,display=True),fill=b.rgb(b.NAVY))
            d.text((x+27,y+62),desc,font=b.font(18),fill=b.rgb(b.MUTED))
    return save(im,"02_superficies.png")


def cockpit():
    im,d=b.canvas("O cockpit de campanha já é robusto","A página /dashboard/campaign/:id é a casa natural do pós-lançamento.")
    b.rounded(d,(90,215,1120,910),35,b.NAVY)
    d.text((140,260),"DASHBOARD DA CAMPANHA",font=b.font(36,display=True),fill=b.rgb(b.WHITE))
    sections=[
        ("ECONOMIA",["Investimento","Revenue","Lucro","ROAS","ROI + imposto"],b.CYAN),
        ("MÍDIA",["Impressões","Cliques","CTR","CPC","Conversões","Custo/conversão"],b.ORANGE),
        ("MONETIZAÇÃO",["GAM requests","Fill rate","eCPM","Viewability","Receita por placement"],b.PURPLE),
        ("DECISÃO",["Orientação","Otimização","Ação de lance","Negativação","URLs do funil"],b.GREEN),
    ]
    for i,(tag,items,color) in enumerate(sections):
        col=i%2; row=i//2; x=140+col*475; y=345+row*260
        b.rounded(d,(x,y,x+430,y+210),24,"151E35",color,2)
        d.text((x+22,y+20),tag,font=b.font(20,bold=True),fill=b.rgb(color))
        for j,item in enumerate(items):
            xx=x+22+(j%2)*195; yy=y+70+(j//2)*42
            d.ellipse((xx,yy+4,xx+9,yy+13),fill=b.rgb(color)); d.text((xx+18,yy),item,font=b.font(18),fill=b.rgb(b.WHITE))
    b.arrow(d,(1150,560),(1325,560),b.ORANGE,8,24)
    b.rounded(d,(1340,300,1710,820),35,b.PALE_ORANGE,b.ORANGE,3)
    d.text((1380,350),"O GAP REAL",font=b.font(23,bold=True),fill=b.rgb(b.ORANGE))
    b.text_wrap(d,"A nova campanha nasce no Hub de Tráfego, mas sua origem, prova, política, funil e recibo ainda não desembocam nesta casa como uma única história.",(1380,415),290,b.font(25),b.INK,9)
    d.text((1380,690),"PRIORIDADE",font=b.font(19,bold=True),fill=b.rgb(b.ORANGE))
    b.text_wrap(d,"Construir a ponte. Não outro painel.",(1380,730),290,b.font(27,display=True),b.NAVY,7)
    return save(im,"03_cockpit_existente.png")


def data_snapshot():
    im,d=b.canvas("Pulso do banco em 22/08/2026","Contagens exatas via PostgREST. Datas indicam o último registro observado.")
    rows=[
        ("Comportamento","raw_events","923","22/08",b.GREEN),
        ("Página","fact_page_daily","671","22/08",b.GREEN),
        ("Projeto","daily_project_metrics","21","22/08",b.GREEN),
        ("Campanha","daily_campaign_metrics","96","19/08",b.ORANGE),
        ("Visitas","site_visits","9.468","12/08",b.ORANGE),
        ("AdSense","adsense_metrics","147","19/02",b.PURPLE),
        ("GAM","gam_metrics","0","—",b.MUTED),
        ("JoinAds","joinads_metrics","0","—",b.MUTED),
        ("Conversões","conversion_queue / batches","0 / 0","—",b.MUTED),
        ("Campanhas","campaigns","4","19/08",b.CYAN),
    ]
    d.text((105,215),"CAPACIDADE",font=b.font(18,bold=True),fill=b.rgb(b.MUTED));d.text((445,215),"OBJETO",font=b.font(18,bold=True),fill=b.rgb(b.MUTED));d.text((1100,215),"LINHAS",font=b.font(18,bold=True),fill=b.rgb(b.MUTED));d.text((1390,215),"ÚLTIMO",font=b.font(18,bold=True),fill=b.rgb(b.MUTED))
    for i,(cap,obj,count,last,color) in enumerate(rows):
        y=260+i*67
        d.rectangle((90,y,1710,y+55),fill=b.rgb(b.WHITE if i%2==0 else "F8FAFC"))
        d.rectangle((90,y,100,y+55),fill=b.rgb(color))
        d.text((120,y+15),cap,font=b.font(21,bold=True),fill=b.rgb(b.NAVY))
        d.text((445,y+16),obj,font=b.font(19),fill=b.rgb(b.INK))
        d.text((1120,y+14),count,font=b.font(22,bold=True),fill=b.rgb(color))
        d.text((1390,y+16),last,font=b.font(20),fill=b.rgb(b.MUTED))
    d.text((100,970),"Leitura: o sistema tem pulso, mas as fontes não compartilham o mesmo frescor.",font=b.font(23,display=True),fill=b.rgb(b.ORANGE))
    return save(im,"04_pulso_banco.png")


def truth_split():
    im,d=b.canvas("A divergência que contamina as demais","Estado do inventário n8n: ativo declarado não significa dado chegando ao banco do produto.",dark=True)
    b.rounded(d,(100,300,500,720),35,"151E35",b.CYAN,3)
    d.text((160,350),"PRODUTO",font=b.font(30,display=True),fill=b.rgb(b.WHITE));d.text((160,410),"Front + Backend",font=b.font(23),fill=b.rgb(b.CYAN));b.text_wrap(d,"Lê e escreve no Supabase self-hosted.",(160,485),280,b.font(25),b.WHITE,8)
    b.rounded(d,(700,240,1110,780),35,"151E35",b.YELLOW,3)
    d.text((765,295),"n8n",font=b.font(42,display=True),fill=b.rgb(b.WHITE));d.text((765,365),"30 workflows",font=b.font(26),fill=b.rgb(b.YELLOW));d.text((765,415),"23 ativos declarados",font=b.font(24),fill=b.rgb(b.YELLOW));d.text((765,485),"271 refs → hospedado",font=b.font(24,bold=True),fill=b.rgb(b.ORANGE));d.text((765,535),"30 refs → self-hosted",font=b.font(24,bold=True),fill=b.rgb(b.CYAN));b.text_wrap(d,"Sem recibo comum, “verde” pode terminar no banco errado.",(765,620),300,b.font(22),b.WHITE,8)
    b.rounded(d,(1310,225,1700,490),35,"151E35",b.ORANGE,3)
    d.text((1370,275),"HOSPEDADO",font=b.font(27,display=True),fill=b.rgb(b.WHITE));b.text_wrap(d,"Banco legado que o produto atual não lê.",(1370,345),260,b.font(23),b.WHITE,8)
    b.rounded(d,(1310,600,1700,865),35,"151E35",b.GREEN,3)
    d.text((1370,650),"SELF-HOSTED",font=b.font(27,display=True),fill=b.rgb(b.WHITE));b.text_wrap(d,"database.agenciavolc.com.br — banco vivo analisado.",(1370,720),260,b.font(23),b.WHITE,8)
    b.arrow(d,(500,510),(690,510),b.CYAN,7,22);b.arrow(d,(1110,390),(1295,360),b.ORANGE,7,22);b.arrow(d,(1110,640),(1295,710),b.CYAN,7,22)
    return save(im,"05_verdade_dividida.png")


def conversion_loop():
    im,d=b.canvas("O elo com maior alavancagem","O sinal já foi capturado. Falta transformá-lo em conversão enviada e reconciliada.")
    steps=[
        (180,"VISITA","9.468",b.GREEN),(485,"IDENTIFICAÇÃO","9.407 GCLIDs",b.GREEN),
        (790,"VALOR","0 calculado",b.ORANGE),(1095,"FILA","0 itens",b.MUTED),
        (1400,"LOTE + ENVIO","0 lotes",b.MUTED),(1640,"GOOGLE","sem sinal",b.MUTED),
    ]
    for i,(x,title,value,color) in enumerate(steps):
        b.rounded(d,(x-120,390,x+120,620),30,b.WHITE,color,4)
        d.text((x-88,430),title,font=b.font(20,bold=True),fill=b.rgb(color))
        tw=d.textbbox((0,0),value,font=b.font(28,display=True))[2];d.text((x-tw/2,510),value,font=b.font(28,display=True),fill=b.rgb(b.NAVY))
        if i<len(steps)-1:b.arrow(d,(x+122,505),(steps[i+1][0]-122,505),"AAB5C5",5,16)
    b.rounded(d,(350,760,1450,900),30,b.PALE_ORANGE,b.ORANGE,3)
    d.text((410,800),"Não é um novo sistema de monitoramento.",font=b.font(28,display=True),fill=b.rgb(b.ORANGE))
    b.text_wrap(d,"É fechar o circuito que ensina o Google Ads sobre a qualidade econômica do clique.",(410,850),940,b.font(23),b.INK,6)
    return save(im,"06_loop_conversao.png")


def traffic_bridge():
    im,d=b.canvas("Curto prazo: terminar Tráfego sem duplicar o sistema","A entrega é uma ponte entre dois produtos que já existem.")
    b.rounded(d,(90,265,700,800),36,b.PALE_BLUE,b.CYAN,3)
    d.text((145,315),"NOVA CAMPANHA SEARCH",font=b.font(30,display=True),fill=b.rgb(b.NAVY))
    for i,t in enumerate(["Origem e funil","Keywords e negativos","Copy persistida","Conta e escopo","Mesa de lance","Prova real","Criação pausada + recibo"]):
        y=390+i*50; d.ellipse((145,y+4,156,y+15),fill=b.rgb(b.CYAN));d.text((172,y),t,font=b.font(22),fill=b.rgb(b.INK))
    b.arrow(d,(735,535),(1045,535),b.ORANGE,12,32)
    d.text((785,460),"PONTE",font=b.font(25,bold=True),fill=b.rgb(b.ORANGE))
    for i,t in enumerate(["campaign_id","funnel_run_id","URLs","selo + política","recibo"]):
        d.text((790,510+i*38),t,font=b.font(18),fill=b.rgb(b.MUTED))
    b.rounded(d,(1080,265,1710,800),36,b.NAVY,b.ORANGE,3)
    d.text((1135,315),"COCKPIT EXISTENTE",font=b.font(30,display=True),fill=b.rgb(b.WHITE))
    for i,t in enumerate(["Estado e ativação","Custo e entrega","Receita e lucro","ROAS / ROI","Conversões","GAM / Display","Orientação e ação","Funil e comportamento"]):
        y=390+i*46; d.ellipse((1135,y+4,1146,y+15),fill=b.rgb(b.ORANGE));d.text((1162,y),t,font=b.font(21),fill=b.rgb(b.WHITE))
    d.text((560,925),"Prova de saída: lançar → clicar no recibo → operar a mesma campanha no cockpit, com a história preservada.",font=b.font(22,display=True),fill=b.rgb(b.DEEP))
    return save(im,"07_ponte_trafego.png")


def attribution():
    im,d=b.canvas("A cadeia de identidade que falta fechar","Cada etapa existe; os campos ainda não formam uma corrente completa.")
    items=[
        ("OPORTUNIDADE","Pautador",b.CYAN),("RUN","Redator",b.PURPLE),("CAMPANHA","Google Ads",b.ORANGE),
        ("URL","Página",b.CYAN),("HOST + PATH","Sensor",b.GREEN),("RESULTADO","Cockpit",b.DEEP),
    ]
    for i,(title,desc,color) in enumerate(items):
        x=120+i*285
        b.rounded(d,(x,350,x+235,620),30,b.WHITE,color,4)
        d.text((x+24,390),title,font=b.font(20,bold=True),fill=b.rgb(color));d.text((x+24,450),desc,font=b.font(28,display=True),fill=b.rgb(b.NAVY))
        status="OK" if i<2 or i==5 else "PARCIAL"
        b.rounded(d,(x+24,535,x+145,583),17,b.GREEN if status=="OK" else b.ORANGE)
        d.text((x+45,548),status,font=b.font(17,bold=True),fill=b.rgb(b.WHITE))
        if i<len(items)-1:b.arrow(d,(x+237,485),(x+280,485),"AAB5C5",5,15)
    d.text((165,760),"Medição atual",font=b.font(21,bold=True),fill=b.rgb(b.MUTED))
    b.text_wrap(d,"1/4 campanha com funnel_run_id  •  campanha nova com 0 URLs  •  671 fatos de página com host vazio  •  34 fatos com ad views positivos",(165,805),1450,b.font(28,display=True),b.NAVY,9)
    return save(im,"08_atribuicao.png")


def fact_opinion_action():
    im,d=b.canvas("Fato, recomendação e ação são coisas diferentes","Separar os substantivos torna o ORAKUL explicável e a execução auditável.")
    cards=[
        (120,"FATO","O que foi medido","custo · receita · clique · conversão · frescor",b.CYAN),
        (650,"RECOMENDAÇÃO","O que o motor concluiu","engine · versão · razão · confiança · proposta",b.PURPLE),
        (1180,"AÇÃO","O que foi autorizado e executado","dono · parâmetros · recibo · verificação",b.ORANGE),
    ]
    for x,title,sub,desc,color in cards:
        b.rounded(d,(x,285,x+480,760),35,b.WHITE,color,4)
        d.text((x+38,335),title,font=b.font(37,display=True),fill=b.rgb(color));d.text((x+38,415),sub,font=b.font(25),fill=b.rgb(b.NAVY));b.text_wrap(d,desc,(x+38,500),390,b.font(27),b.MUTED,10)
    b.arrow(d,(602,520),(642,520),"AAB5C5",5,15);b.arrow(d,(1132,520),(1172,520),"AAB5C5",5,15)
    d.text((325,895),"Hoje: métricas + orientação + otimização dividem a mesma linha; bid_actions = 0.",font=b.font(26,display=True),fill=b.rgb(b.ORANGE))
    return save(im,"09_fato_opiniao_acao.png")


def priority_stack(graph):
    im,d=b.canvas("Prioridades corrigidas","Ordem por capacidade destravada e dependência — não por novidade visual.")
    colors=[b.ORANGE,b.PURPLE,b.CYAN,b.ORANGE,b.PURPLE,b.DEEP,b.CYAN,b.GREEN]
    for i,p in enumerate(graph["priorities"]):
        col=i%2;row=i//2;x=90+col*850;y=215+row*185;color=colors[i]
        b.rounded(d,(x,y,x+800,y+150),26,b.WHITE,"D8DEE9",2)
        b.rounded(d,(x+22,y+28,x+82,y+88),30,color)
        d.text((x+43,y+41),str(p["rank"]),font=b.font(23,bold=True),fill=b.rgb(b.WHITE))
        d.text((x+108,y+27),p["title"],font=b.font(24,display=True),fill=b.rgb(b.NAVY))
        b.text_wrap(d,p["why"],(x+108,y+70),650,b.font(18),b.MUTED,6,max_lines=2)
    return save(im,"10_prioridades.png")


def workstreams():
    im,d=b.canvas("Três trilhos, uma prioridade ativa por vez","O sistema pode preservar o todo sem misturar filas de execução.")
    lanes=[
        ("PLATAFORMA VOLC O.S.","verdade → conversão → ponte de tráfego → dados → governança",b.CYAN),
        ("OPERAÇÃO DE MÍDIA","campanhas existentes → leitura → decisão humana → atuação segura",b.ORANGE),
        ("PUBLISHER / FOCO GENIAL","CWV → ad loader → dataLayer → correlação receita × experiência",b.PURPLE),
    ]
    for i,(title,desc,color) in enumerate(lanes):
        y=260+i*235
        b.rounded(d,(110,y,1690,y+180),30,b.WHITE,color,4)
        d.text((155,y+35),title,font=b.font(28,display=True),fill=b.rgb(color));b.text_wrap(d,desc,(155,y+92),1360,b.font(26),b.NAVY,8)
    d.text((395,950),"ClickUp Foco Genial: 11 tarefas oficiais, todas ainda em “to do”. Elas não substituem o roadmap do produto.",font=b.font(22,display=True),fill=b.rgb(b.DEEP))
    return save(im,"11_trilhos.png")


def build(graph, images):
    doc=Document();b.style_doc(doc);sec=doc.sections[0];b.configure_section(sec,cover=True)
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(0);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(images["cover"]),width=Cm(21),height=Cm(29.7))
    sec=doc.add_section(WD_SECTION.NEW_PAGE);b.configure_section(sec)

    b.add_title(doc,"O que mudou nesta versão","A versão 0.2 começa reconhecendo uma leitura incorreta da versão anterior.","Nota de correção")
    b.callout(doc,"Correção principal","O VOLC O.S. já possui monitoramento robusto por campanha em /dashboard/campaign/:campaignId. Portanto, “construir monitoramento” não é uma prioridade correta. A prioridade é ligar a nova criação Search ao cockpit existente e garantir que os dados que o alimentam sejam atuais, reconciliados e rastreáveis.",b.ORANGE,b.PALE_ORANGE)
    doc.add_heading("O que passou a ser a fonte de verdade",level=2)
    b.simple_table(doc,["CAMADA","MEDIÇÃO DESTA EDIÇÃO"],[
        ("Produto","22 rotas/telas ligadas a serviços e objetos de dados."),
        ("Banco","64 tabelas/views, 67 funções e contagens exatas via PostgREST."),
        ("Automação","30 workflows n8n, 2.311 nós e 23 ativos declarados no inventário."),
        ("Backlog publisher","11 tarefas ClickUp oficiais do Foco Genial."),
        ("Grafo","269 nós e 442 relações em um artefato navegável e legível por máquina."),
    ],[5,11.9])
    b.callout(doc,"Status do documento anterior","A v0.1 permanece como registro do processo, mas sua ordem de prioridades está superada por esta edição.",b.PURPLE,b.PALE_PURPLE)

    b.page_break(doc);b.add_title(doc,"Como usar este workbook","Abra o Mapa Mestre para explorar; use este documento para decidir e executar.","Orientação")
    b.simple_table(doc,["ARTEFATO","FUNÇÃO"],[
        ("Mapa_Mestre_VOLC_OS.html","Explorar capacidades, relações, estados, evidências e prioridades."),
        ("volc-os-graph.json","Fonte de verdade para futuras análises, documentação e automações."),
        ("Este workbook v0.2","Narrativa executiva, ordem de trabalho, decisões e provas de saída."),
        ("ClickUp","Receber somente o trabalho aprovado e executável de cada trilho."),
    ],[6,10.9])
    b.callout(doc,"Regra permanente","Nenhuma nova prioridade entra no workbook sem primeiro ganhar um nó no grafo, relações, estado, evidência e pré-requisitos.",b.DEEP,b.PALE_BLUE)

    b.page_break(doc);b.add_section_divider(doc,"01","O todo antes da fila","O sistema já contém aquisição, produção, monetização, comportamento, decisão e governança. O desafio é fechar as relações entre essas partes.",b.CYAN)
    b.page_break(doc);b.add_title(doc,"Mapa executivo do sistema","A visão executiva mostra capacidades; o inventário navegável revela telas, tabelas, RPCs, workflows e documentos.","Mapa Mestre")
    b.add_image(doc,images["graph"]);b.add_caption(doc,"Figura 1 — 24 capacidades distribuídas em 10 clusters. Fonte: Mapa Mestre VOLC O.S., 22/08/2026.")
    b.callout(doc,"Tamanho medido","269 nós e 442 relações: 24 capacidades, 22 superfícies, 64 tabelas/views, 67 funções de banco, 30 workflows, 16 sistemas externos, 14 serviços, 11 tarefas ClickUp e outros módulos/documentos.",b.CYAN,b.PALE_BLUE)

    b.page_break(doc);b.add_title(doc,"O produto que já existe","O VOLC O.S. já é um sistema operacional amplo, não uma camada recém-iniciada de tráfego.","Superfícies atuais")
    b.add_image(doc,images["surfaces"]);b.add_caption(doc,"Figura 2 — Superfícies registradas no aplicativo atual.")
    b.simple_table(doc,["JORNADA","O QUE JÁ ENTREGA"],[
        ("Operar","Portfólio, projeto, campanha, relatórios, economia e configuração."),
        ("Produzir","Descoberta, mineração, funil, páginas, prova e publicação."),
        ("Adquirir","Fila, copy, keywords, lance, prova e criação Search protegida."),
        ("Decidir","Orientação, otimização, ação de lance e Display dentro do cockpit."),
        ("Governar","Usuários, custos, integrações, permissões e nova camada RBAC em evolução."),
    ],[4,12.9])

    b.page_break(doc);b.add_title(doc,"O cockpit de campanha já é a casa do pós-lançamento","A página apontada pelo dono muda completamente o diagnóstico.","Correção central")
    b.add_image(doc,images["cockpit"]);b.add_caption(doc,"Figura 3 — Capacidades observadas em CampaignDetailDashboard.tsx.")
    b.callout(doc,"Nova tese para a Camada de Tráfego","Ela não precisa inventar Observação, Diagnóstico e Decisão. Precisa entregar a campanha recém-criada para a estrutura que já faz isso, preservando identidade e evidências.",b.ORANGE,b.PALE_ORANGE)

    b.page_break(doc);b.add_section_divider(doc,"02","Verdade operacional","O problema não é ausência de dados. É diferença de frescor, origem e continuidade entre fontes.",b.ORANGE)
    b.page_break(doc);b.add_title(doc,"Pulso atual do Supabase","O banco self-hosted contém história real da operação e também mostra exatamente onde a corrente para.","Medição viva")
    b.add_image(doc,images["snapshot"]);b.add_caption(doc,"Figura 4 — Contagens exatas e última data observada em 22/08/2026.")
    b.callout(doc,"Leitura correta","Comportamento, fatos de página e métricas de projeto chegaram a 22/08. Métricas de campanha terminaram em 19/08. AdSense é histórico; GAM e JoinAds estão vazios no banco vivo. Portanto, o cockpit existe, mas sua verdade depende de pipelines com frescor desigual.",b.ORANGE,b.PALE_ORANGE)

    b.page_break(doc);b.add_title(doc,"A divisão que contamina o restante","O app e os workflows ainda não compartilham um destino operacional único.","Fonte de verdade")
    b.add_image(doc,images["truth"]);b.add_caption(doc,"Figura 5 — Referências medidas no inventário sanitizado do n8n.")
    b.simple_table(doc,["FATO","IMPLICAÇÃO"],[
        ("Produto lê o self-hosted","Esse banco é a verdade visível ao operador."),
        ("271 referências n8n apontam para o hospedado","Rotina ativa pode trabalhar sem alimentar o produto."),
        ("30 referências apontam para o self-hosted","Somente uma parte do parque converge para a mesma verdade."),
        ("“Ativo” é estado declarado","Sem recibo, não prova que dados chegaram ou foram aceitos."),
    ],[7,9.9])
    b.callout(doc,"Prioridade global nº 1","Decidir e executar a unificação da verdade operacional antes de ampliar automações.",b.ORANGE,b.PALE_ORANGE)

    b.page_break(doc);b.add_title(doc,"O loop de conversão offline continua aberto","Aqui está a maior capacidade realmente ausente — e não uma tela.","Alavancagem econômica")
    b.add_image(doc,images["conversion"]);b.add_caption(doc,"Figura 6 — Estado medido do circuito de conversão em 22/08/2026.")
    b.simple_table(doc,["MEDIÇÃO","RESULTADO"],[
        ("Visitas registradas","9.468"),("Visitas com GCLID","9.407"),("Status pending","9.468"),
        ("Valor calculado","0"),("conversion_queue","0"),("conversion_batches","0"),
    ],[8,8.9])
    b.callout(doc,"Por que vem antes de ORAKUL","O Google não pode aprender a qualidade econômica do tráfego se o sinal que já capturamos nunca volta para a plataforma.",b.PURPLE,b.PALE_PURPLE)

    b.page_break(doc);b.add_title(doc,"Atribuição é a corrente central","O sistema conhece as peças, mas ainda perde a identidade entre campanha, funil e comportamento.","Elo de negócio")
    b.add_image(doc,images["attribution"]);b.add_caption(doc,"Figura 7 — Cadeia de correspondências e medições atuais.")
    b.simple_table(doc,["ELO","ESTADO MEDIDO","FECHAMENTO"],[
        ("Run → campanha","1 de 4 campanhas possui funnel_run_id.","Persistir em toda criação nova."),
        ("Campanha → URLs","Campanha 24155134757 possui 0 URLs associadas.","Copiar URLs publicadas no lançamento."),
        ("Página → site","671 fatos de página com host vazio.","Preservar host na agregação."),
        ("Página → anúncios","34 fatos têm ad_views > 0; 24/33 funis têm média zero.","Completar sensor e regra de agregação."),
    ],[4,6.5,6.4],font_size=8.2)

    b.page_break(doc);b.add_section_divider(doc,"03","Prioridades corrigidas","Primeiro verdade e sinal; depois integração, reconciliação, atribuição e governança; só então inteligência avançada.",b.PURPLE)
    b.page_break(doc);b.add_title(doc,"A nova ordem","O curto prazo da Camada de Tráfego continua importante — agora ocupa o lugar correto dentro do sistema.","Roadmap global")
    b.add_image(doc,images["priorities"]);b.add_caption(doc,"Figura 8 — Ordem qualitativa baseada em dependências e evidência atual.")
    for p in graph["priorities"]:
        if p["rank"] == 5:
            b.page_break(doc)
            b.add_title(doc,"Prioridades 5–8","Fechar o contexto operacional antes de reabrir inteligência e expansão.","Continuação do roadmap")
        b.callout(doc,f"{p['rank']} · {p['title']}",p["proof"],b.DEEP if p["rank"]>3 else b.ORANGE,b.PALE_BLUE if p["rank"]>3 else b.PALE_ORANGE)

    b.page_break(doc);b.add_title(doc,"Prioridade 3 — a ponte de Tráfego","Este é o pacote de curto prazo que termina a experiência sem duplicar o cockpit.","Camada de Tráfego")
    b.add_image(doc,images["bridge"]);b.add_caption(doc,"Figura 9 — A jornada corrigida: nascimento no Tráfego, operação no cockpit existente.")
    b.simple_table(doc,["ENTREGA","PRONTO QUANDO"],[
        ("Recibo navegável","Após subir, a ação principal abre a campanha criada no cockpit existente."),
        ("Identidade única","campaign_id, project_id, opportunity_id e funnel_run_id permanecem ligados."),
        ("URLs automáticas","As páginas publicadas entram em campaign_funnel_urls sem digitação manual."),
        ("Prova preservada","Selo, política, conta, estratégia e payload ficam consultáveis na história."),
        ("Retorno consistente","Quadro, sino e campanha no ar apontam para a mesma casa da campanha."),
        ("Sem tela paralela","Métricas e diagnósticos continuam no dashboard já existente."),
    ],[6.2,10.7])
    b.callout(doc,"Prova de saída","Criar uma Search no Hub → abrir o recibo → chegar ao dashboard correto → entender origem, configuração e estado sem reconstruir nada.",b.GREEN,b.PALE_GREEN)

    b.page_break(doc);b.add_title(doc,"Prioridade 4 — frescor e reconciliação","Monitoramento existe. O que precisa melhorar é a confiabilidade da alimentação.","Dados operacionais")
    b.simple_table(doc,["FONTE","HOJE","PROVA DE FECHAMENTO"],[
        ("Google Ads / campanha","Último dado 19/08.","D0 e D-1 chegam ao banco certo, com recibo e divergência."),
        ("Projeto / receita","Chegou a 22/08.","Moeda e fonte explícitas; projeto reconcilia com campanha/URL."),
        ("AdSense","Histórico até 19/02.","Fonte reativada ou aposentada explicitamente."),
        ("GAM","Tabela vazia.","Primeiro lote aceito e reconciliado com requests/receita."),
        ("JoinAds","Tabela vazia no snapshot.","Carga chega ao banco vivo e deixa recibo por janela."),
        ("Câmbio","Último efetivo 18/02.","Uma taxa por data e conversão reprocessável."),
    ],[4.1,5.1,7.7],font_size=8.0)
    b.callout(doc,"Mudança de linguagem","Não “construir monitoramento”; restabelecer frescor, proveniência e reconciliação das fontes que o monitoramento já apresenta.",b.ORANGE,b.PALE_ORANGE)

    b.page_break(doc);b.add_title(doc,"Prioridade 6 — governar decisão e atuação","A visão avançada existe; falta dar substantivos comuns e rastreáveis para ela.","ORAKUL moderno")
    b.add_image(doc,images["fao"]);b.add_caption(doc,"Figura 10 — Separação recomendada do domínio operacional.")
    b.simple_table(doc,["OBJETO","CONTEÚDO MÍNIMO","REGRA"],[
        ("Fato","fonte, janela, moeda, valor, frescor, qualidade","Nunca é reescrito por uma opinião."),
        ("Recomendação","engine, versão, razão, confiança, parâmetros","Não executa sozinha."),
        ("Autorização","dono, escopo, limite, validade","Aumento e mudança de rota exigem humano."),
        ("Ação","o que foi enviado, resposta, recibo, verificação","Pode ser ligada ao fato e à recomendação."),
    ],[3.1,7.7,6.1])
    b.callout(doc,"Estado medido","12 orientações reais e 10 otimizações históricas convivem em daily_campaign_metrics; bid_actions continua com 0 linhas.",b.PURPLE,b.PALE_PURPLE)

    b.page_break(doc);b.add_title(doc,"Prioridade 7 — simplificar as rotinas","O objetivo não é trocar n8n por princípio. É tornar cada rotina confiável, única e observável.","Orquestração")
    b.simple_table(doc,["MEDIÇÃO","LEITURA"],[
        ("30 workflows / 2.311 nós","O parque tem cobertura grande e alto custo cognitivo."),
        ("23 ativos declarados","Ativo não prova sucesso nem destino correto."),
        ("13 workflows às 06:00 no inventário anterior","Há concorrência, duplicação e risco de quota."),
        ("Pares D0/D-1","O conceito de provisório + fechado é bom e deve ser preservado."),
        ("pg_cron comportamental","É o contraexemplo simples: data, idempotência e histórico."),
    ],[6,10.9])
    doc.add_heading("Contrato recomendado por rotina",level=2)
    for item in ["Uma fonte e um grão explícitos.","Uma janela de data parametrizável.","Idempotência pela chave natural.","Recibo com linhas pedidas, aceitas, descartadas e atualizadas.","Frescor publicado no mesmo banco que o produto lê.","Alerta acionável e owner.","Modo sob demanda para reconciliação."]:
        b.bullet(doc,item,checkbox=True)

    b.page_break(doc);b.add_title(doc,"Três trilhos sem competição indevida","Produto, operação de mídia e publisher avançam juntos, mas não compartilham a mesma fila ativa.","Organização")
    b.add_image(doc,images["workstreams"]);b.add_caption(doc,"Figura 11 — Separação de trilhos do roadmap.")
    b.callout(doc,"ClickUp incorporado","A lista oficial contém 11 tarefas: 8 P0 e 3 P1 pelo nome, todas em “to do”. O conteúdo cobre CWV, ad loader, dataLayer, LCP, INP, cache, lazy loading e correlação com receita.",b.PURPLE,b.PALE_PURPLE)

    b.page_break(doc);b.add_title(doc,"O que continua estacionado","Essas capacidades pertencem ao mapa; não pertencem à fila ativa.","Horizonte avançado")
    b.simple_table(doc,["CAPACIDADE","POR QUE ESPERA","GATILHO PARA REABRIR"],[
        ("Preditivo / Crystal Ball","Sem tabela de previsão e histórico reconciliado.","Série confiável + replay + erro medido."),
        ("ORAKUL autônomo","Fato, recomendação e ação ainda misturados.","Proposta governada + executor auditável."),
        ("PMax / Demand Gen","Search ainda precisa fechar sua ponte operacional.","Contrato de canal comprovado."),
        ("Display automatizado","UI existe, mas dados de placement estão vazios.","Custo e receita por placement reconciliados."),
        ("Rede massiva de personas","Arquitetura histórica sem prova operacional atual.","Caso de uso e medição definidos."),
        ("Pricing GAM avançado","Monetização por placement ainda não flui.","Inventário e economia confiáveis."),
    ],[4.3,6.5,6.1],font_size=8.1)

    b.page_break(doc);b.add_section_divider(doc,"04","Execução sem se perder","Toda tarefa precisa apontar para um nó, uma relação e uma prova de saída.",b.GREEN)
    b.page_break(doc);b.add_title(doc,"Quadro de decisões do dono","Estas escolhas mudam a arquitetura ou o risco e não devem ficar escondidas em tarefas.","Decisões abertas")
    b.simple_table(doc,["DECISÃO","PERGUNTA","RESPOSTA"],[
        ("Banco oficial","O self-hosted é confirmado como única verdade operacional?","________________________"),
        ("Workflows legados","Quem autoriza repontar/desligar rotinas do hospedado?","________________________"),
        ("Conversão oficial","Qual ação/label e qual valor entram no primeiro lote?","________________________"),
        ("Ativação Search","Quem ativa e qual portão econômico mínimo?","________________________"),
        ("Ação automática","Quais ações defensivas podem executar sem aprovação?","________________________"),
        ("Moeda JoinAds","Qual moeda a fonte efetivamente retorna?","________________________"),
        ("Trilho ativo","Plataforma, mídia ou publisher é o trilho principal deste ciclo?","________________________"),
    ],[4,8.2,4.7],font_size=8.2)
    b.callout(doc,"Decisão recomendada para começar","Confirmar a fonte única de verdade e congelar novas automações até cada rotina crítica declarar destino e recibo.",b.ORANGE,b.PALE_ORANGE)

    b.page_break(doc);b.add_title(doc,"Ficha de iniciativa ligada ao grafo","O campo “nós afetados” impede tarefas isoladas do sistema.","Template")
    fields=[
        ("INICIATIVA","________________________________________________________"),
        ("TRILHO","☐ Plataforma  ☐ Operação de mídia  ☐ Publisher"),
        ("NÓS AFETADOS NO GRAFO","________________________________________________________"),
        ("RELAÇÃO QUE SERÁ CRIADA OU CORRIGIDA","________________________________________________________"),
        ("ESTADO ATUAL + EVIDÊNCIA","________________________________________________________\n________________________________________________________"),
        ("DECISÕES / PRÉ-REQUISITOS","________________________________________________________"),
        ("FORA DO ESCOPO","________________________________________________________"),
        ("PROVA DE SAÍDA","________________________________________________________\n________________________________________________________"),
        ("DONO / REVISÃO","________________________________________________________"),
    ]
    t=doc.add_table(rows=0,cols=1);t.autofit=False;t.columns[0].width=Cm(17.2)
    for label,value in fields:
        c=t.add_row().cells[0];b.set_cell_shading(c,b.WHITE);b.set_cell_margins(c,145,180,170,180);b.set_cell_border(c,bottom={"val":"single","sz":"8","color":"DDE4EE"})
        p=c.paragraphs[0];r=p.add_run(label+"\n");r.bold=True;r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string(b.DEEP);r=p.add_run(value);r.font.size=Pt(10);r.font.color.rgb=RGBColor.from_string(b.MUTED)

    b.page_break(doc);b.add_title(doc,"Plano da próxima sessão","A próxima conversa pode começar por decisão, não por nova coleta desordenada.","Partida")
    steps=[
        ("Abrir o Mapa Mestre","Confirmar se os 10 clusters representam o sistema como você o enxerga."),
        ("Confirmar a verdade operacional","Decidir banco oficial e destino das rotinas legadas."),
        ("Aprovar a ordem 1–8","Mover apenas divergências reais, com evidência."),
        ("Fechar o pacote da ponte de Tráfego","Transformar a prioridade 3 em iniciativas pequenas e prováveis."),
        ("Levar ao ClickUp por trilho","Produto e publisher deixam de competir na mesma lista mental."),
    ]
    for i,(title,desc) in enumerate(steps,1):
        table=doc.add_table(rows=1,cols=2);table.autofit=False;table.columns[0].width=Cm(1.3);table.columns[1].width=Cm(15.8)
        b.set_cell_shading(table.cell(0,0),b.DEEP);b.set_cell_shading(table.cell(0,1),b.WHITE)
        for c in table.rows[0].cells:b.set_cell_margins(c,150,170,150,170);b.set_cell_border(c,bottom={"val":"single","sz":"7","color":"DDE4EE"})
        p=table.cell(0,0).paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(str(i));r.bold=True;r.font.size=Pt(16);r.font.color.rgb=RGBColor.from_string(b.WHITE)
        p=table.cell(0,1).paragraphs[0];r=p.add_run(title+"\n");r.bold=True;r.font.name="Space Grotesk";r.font.size=Pt(12);r.font.color.rgb=RGBColor.from_string(b.NAVY);r=p.add_run(desc);r.font.size=Pt(9.5);r.font.color.rgb=RGBColor.from_string(b.MUTED)
    b.callout(doc,"Resultado esperado","Sair com uma fonte de verdade confirmada, um trilho principal e o pacote exato da ponte Tráfego → Cockpit.",b.GREEN,b.PALE_GREEN)

    b.page_break(doc);b.add_title(doc,"Evidências desta edição","Números medidos e limites explícitos.","Rastreabilidade")
    b.simple_table(doc,["MEDIÇÃO","RESULTADO","FONTE"],[
        ("Grafo","269 nós, 442 relações","Código + snapshots + inventários"),
        ("Supabase","64 tabelas/views, 67 RPCs","PostgREST vivo, count exact"),
        ("Produto","22 superfícies registradas","src/App.tsx"),
        ("n8n","30 workflows, 2.311 nós, 23 ativos declarados","inventario-n8n/*.meta.json"),
        ("Referências de banco no n8n","271 hospedado / 30 self-hosted","Busca nos workflows sanitizados"),
        ("Tracking","9.468 visitas; 9.407 GCLIDs; 9.468 pending","Supabase, 22/08/2026"),
        ("Cockpit","Métricas, monetização, orientação, otimização e ações","CampaignDetailDashboard.tsx"),
        ("ClickUp Foco Genial","11 tarefas, todas to do","Lista 901328196164"),
    ],[5.7,5.8,5.4],font_size=8.0)
    b.callout(doc,"Limite importante","O estado “ativo” de workflows vem do inventário do n8n e não prova execução bem-sucedida. Esta edição não realizou mutações em banco, ClickUp, n8n ou contas de mídia.",b.ORANGE,b.PALE_ORANGE)

    b.page_break(doc);b.add_title(doc,"Fontes oficiais","A hierarquia agora é mediada pelo grafo, não por um documento isolado.","Apêndice")
    b.simple_table(doc,["FONTE","PAPEL"],[
        ("Sistema e banco vivos","Verdade atual e evidência."),
        ("Mapa Mestre VOLC O.S.","Correspondências e estado de cada capacidade."),
        ("Plano oficial Foco Genial","Estratégia inicial de performance e monetização."),
        ("PRD + SPEC Arbitragem","Hipóteses de evolução e princípios de governança."),
        ("Publisher Global Blueprint","Norte empresarial e subsistemas."),
        ("Second Brain","Memória de modelos, experimentos e ambições."),
        ("Inventário n8n","Genealogia, rotinas, riscos e destinos."),
        ("ClickUp","Backlog operacional do publisher."),
    ],[7,9.9])
    b.callout(doc,"Regra para a v0.3","Novas fontes entram no grafo primeiro. O workbook é regenerado depois, preservando rastreabilidade e evitando outra leitura parcial.",b.DEEP,b.PALE_BLUE)

    b.page_break(doc)
    table=doc.add_table(rows=1,cols=1);table.autofit=False;table.columns[0].width=Cm(17.3)
    c=table.cell(0,0);b.set_cell_shading(c,b.NAVY);b.set_cell_margins(c,900,650,900,650)
    p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Primeiro o mapa.\n");r.font.name="Space Grotesk";r.font.size=Pt(30);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(b.WHITE)
    r=p.add_run("Depois, o próximo passo certo.\n\n");r.font.name="Space Grotesk";r.font.size=Pt(30);r.font.bold=True;r.font.color.rgb=RGBColor.from_string(b.CYAN)
    r=p.add_run("VOLC O.S.  •  Mapa Mestre e Prioridades  •  v0.2");r.font.name="Inter";r.font.size=Pt(10);r.font.color.rgb=RGBColor.from_string("B7C3D9")
    doc.core_properties.title="VOLC O.S. — Mapa Mestre e Prioridades v0.2";doc.core_properties.subject="Cartografia global do sistema e prioridades corrigidas";doc.core_properties.author="VOLC"
    doc.save(OUT)


def main():
    ASSET_DIR.mkdir(parents=True,exist_ok=True)
    graph=json.loads(GRAPH_PATH.read_text())
    images={"cover":cover(),"graph":graph_overview(graph),"surfaces":product_surfaces(),"cockpit":cockpit(),"snapshot":data_snapshot(),"truth":truth_split(),"conversion":conversion_loop(),"bridge":traffic_bridge(),"attribution":attribution(),"fao":fact_opinion_action(),"priorities":priority_stack(graph),"workstreams":workstreams()}
    build(graph,images);print(OUT)


if __name__=="__main__":main()
