#!/usr/bin/env python3
"""Gera o workbook executivo do VOLC O.S. a partir das fontes oficiais analisadas."""

from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "workbook-volc-os" / "assets"
OUT_DIR = ROOT / "entregaveis"
OUT_PATH = OUT_DIR / "Workbook_VOLC_OS_Proximos_Passos_v0.1.docx"

W, H = 1800, 1040
NAVY = "0B1020"
DEEP = "0D47A1"
CYAN = "00D4FF"
PURPLE = "8A2BE2"
ORANGE = "FF3D00"
OFF = "F3F4F6"
WHITE = "FFFFFF"
INK = "1A1C1E"
MUTED = "667085"
PALE_BLUE = "EAF6FF"
PALE_ORANGE = "FFF0EB"
PALE_PURPLE = "F3ECFF"
PALE_GREEN = "EAF8F1"
GREEN = "16845B"
YELLOW = "F6C344"
RED = "D92D20"

SPACE_FONT = Path("/Users/mac/Library/Fonts/SpaceGrotesk-VariableFont_wght.ttf")
BODY_FONT = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
BODY_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
LOGO_LIGHT = ROOT / "public" / "volc-logo-light.png"
V_MARK = ROOT / "public" / "volc-v.png"


def rgb(hex_: str) -> tuple[int, int, int]:
    return tuple(int(hex_[i:i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False, display: bool = False):
    path = SPACE_FONT if display and SPACE_FONT.exists() else (BODY_BOLD if bold else BODY_FONT)
    return ImageFont.truetype(str(path), size=size)


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=rgb(fill), outline=rgb(outline) if outline else None, width=width)


def text_wrap(draw, text, xy, max_width, fnt, fill=INK, line_gap=8, max_lines=None):
    words = text.split()
    lines, current = [], ""
    for word in words:
        candidate = (current + " " + word).strip()
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = lines[-1].rstrip(".,;:") + "…"
    x, y = xy
    ascent = fnt.getbbox("Ag")[3]
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=rgb(fill))
        y += ascent + line_gap
    return y


def arrow(draw, start, end, color=DEEP, width=6, head=16):
    draw.line([start, end], fill=rgb(color), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    pts = [
        end,
        (end[0] - head * math.cos(angle - 0.5), end[1] - head * math.sin(angle - 0.5)),
        (end[0] - head * math.cos(angle + 0.5), end[1] - head * math.sin(angle + 0.5)),
    ]
    draw.polygon(pts, fill=rgb(color))


def canvas(title=None, subtitle=None, dark=False):
    bg = NAVY if dark else OFF
    im = Image.new("RGB", (W, H), rgb(bg))
    d = ImageDraw.Draw(im)
    if title:
        d.text((90, 66), title, font=font(54, display=True), fill=rgb(WHITE if dark else NAVY))
    if subtitle:
        text_wrap(d, subtitle, (92, 135), 1550, font(25), WHITE if dark else MUTED, 6)
    return im, d


def save(im, name):
    path = ASSET_DIR / name
    im.save(path, quality=95)
    return path


def draw_cover():
    im = Image.new("RGB", (1400, 1980), rgb(NAVY))
    d = ImageDraw.Draw(im)
    # Aurora de marca.
    layer = Image.new("RGBA", im.size, (0, 0, 0, 0))
    ld = ImageDraw.Draw(layer)
    for cx, cy, color, radius in [
        (1200, 450, CYAN, 440), (1080, 680, PURPLE, 520), (1260, 950, ORANGE, 380),
        (100, 1770, DEEP, 460),
    ]:
        ld.ellipse((cx-radius, cy-radius, cx+radius, cy+radius), fill=(*rgb(color), 115))
    layer = layer.filter(ImageFilter.GaussianBlur(110))
    im = Image.alpha_composite(im.convert("RGBA"), layer).convert("RGB")
    d = ImageDraw.Draw(im)
    if LOGO_LIGHT.exists():
        logo = Image.open(LOGO_LIGHT).convert("RGBA")
        logo.thumbnail((390, 150))
        im.paste(logo, (90, 90), logo)
    d.text((95, 475), "WORKBOOK", font=font(42, display=True), fill=rgb(CYAN))
    d.text((90, 545), "VOLC O.S.", font=font(115, display=True), fill=rgb(WHITE))
    d.text((95, 685), "Próximos passos", font=font(76, display=True), fill=rgb(WHITE))
    text_wrap(d, "Como terminar a Camada de Tráfego sem perder de vista o sistema inteiro.",
              (100, 810), 1080, font(39), WHITE, 13)
    rounded(d, (95, 1120, 905, 1350), 28, "141D36", CYAN, 3)
    d.text((135, 1163), "FOCO DESTE CICLO", font=font(25, bold=True), fill=rgb(CYAN))
    text_wrap(d, "Transformar o lançamento de campanhas Search em uma operação completa: antes, durante e depois da ativação.",
              (135, 1215), 710, font(30), WHITE, 10)
    d.text((98, 1785), "Versão 0.1  •  22 de agosto de 2026", font=font(22), fill=rgb("B7C3D9"))
    d.text((98, 1830), "Documento vivo — pronto para receber as próximas fontes", font=font(22), fill=rgb("B7C3D9"))
    return save(im, "00_capa.png")


def draw_source_map():
    im, d = canvas("Como as fontes se encaixam", "Cada fonte tem uma função diferente. O código em execução decide o que existe hoje.")
    cards = [
        (90, 220, 390, 520, "NORTE", "Publisher Global Blueprint", "A empresa que queremos construir", CYAN),
        (430, 220, 730, 520, "ESTRATÉGIA", "Plano oficial Foco Genial", "O primeiro caso de uso e suas condições", ORANGE),
        (770, 220, 1070, 520, "ARQUITETURA", "PRD + SPEC Arbitragem", "A hipótese de evolução do produto", PURPLE),
        (1110, 220, 1410, 520, "MEMÓRIA", "Second Brain", "Modelos, testes e ambições anteriores", DEEP),
        (1450, 220, 1710, 520, "ORGANIZAÇÃO", "ClickUp", "Fonte incremental ainda não incorporada", YELLOW),
    ]
    for x1, y1, x2, y2, tag, title, desc, color in cards:
        rounded(d, (x1, y1, x2, y2), 26, WHITE, "D8DEE9", 2)
        d.rectangle((x1, y1, x2, y1+12), fill=rgb(color))
        d.text((x1+24, y1+35), tag, font=font(19, bold=True), fill=rgb(color if color != YELLOW else INK))
        text_wrap(d, title, (x1+24, y1+82), x2-x1-48, font(27, display=True), NAVY, 8)
        text_wrap(d, desc, (x1+24, y1+183), x2-x1-48, font(21), MUTED, 7)
    for x in (240, 580, 920, 1260, 1580):
        arrow(d, (x, 545), (900, 695), "AAB5C5", 4, 14)
    rounded(d, (470, 690, 1330, 900), 32, NAVY)
    d.text((520, 730), "VERDADE ATUAL", font=font(23, bold=True), fill=rgb(CYAN))
    d.text((520, 785), "VOLC O.S. em funcionamento", font=font(42, display=True), fill=rgb(WHITE))
    text_wrap(d, "O que está implementado, testado e observável agora.", (520, 845), 730, font(24), WHITE, 7)
    return save(im, "01_fontes.png")


def draw_global_loop():
    im, d = canvas("A máquina que o VOLC O.S. coordena", "Tráfego é uma engrenagem do negócio — não o negócio inteiro.", dark=True)
    cx, cy, radius = 900, 555, 350
    labels = [
        ("ATENÇÃO", "Demanda e intenção"), ("ATIVOS", "Conteúdo e páginas"),
        ("MONETIZAÇÃO", "Receita por sessão"), ("DADOS", "Aprendizado verificável"),
        ("AUDIÊNCIA PRÓPRIA", "Retorno e LTV"), ("AQUISIÇÃO MELHOR", "Compra mais inteligente"),
    ]
    colors = [CYAN, PURPLE, ORANGE, CYAN, PURPLE, ORANGE]
    points = []
    for i in range(6):
        a = -math.pi/2 + i * 2*math.pi/6
        points.append((cx + radius*math.cos(a), cy + radius*math.sin(a)))
    for i, p in enumerate(points):
        arrow(d, p, points[(i+1) % 6], colors[i], 7, 20)
    for (x, y), (title, desc), color in zip(points, labels, colors):
        rounded(d, (x-145, y-62, x+145, y+62), 22, "151E35", color, 3)
        tw = d.textbbox((0, 0), title, font=font(22, bold=True))[2]
        d.text((x-tw/2, y-32), title, font=font(22, bold=True), fill=rgb(WHITE))
        tw2 = d.textbbox((0, 0), desc, font=font(17))[2]
        d.text((x-tw2/2, y+5), desc, font=font(17), fill=rgb("B7C3D9"))
    rounded(d, (650, 425, 1150, 685), 120, DEEP, CYAN, 4)
    d.text((760, 490), "VOLC O.S.", font=font(48, display=True), fill=rgb(WHITE))
    text_wrap(d, "coordena decisões, execução, prova e aprendizado", (750, 565), 330, font(24), WHITE, 7)
    return save(im, "02_maquina_global.png")


def draw_system_now():
    im, d = canvas("O sistema hoje, em linguagem simples", "Uma cadeia já existe. O principal vazio está no retorno dos resultados para a próxima decisão.")
    xs = [110, 515, 920, 1325]
    stages = [
        ("1", "PAUTADOR", "Encontra e estrutura oportunidades", "PROVADO", GREEN),
        ("2", "REDATOR", "Produz e publica o ativo", "EM EVOLUÇÃO", DEEP),
        ("3", "TRÁFEGO", "Monta, prova e sobe Search pausada", "PROVADO", GREEN),
        ("4", "RESULTADOS", "Devolve custo, receita e aprendizado", "FRAGMENTADO", ORANGE),
    ]
    for x, (n, title, desc, status, color) in zip(xs, stages):
        rounded(d, (x, 285, x+330, 680), 30, WHITE, "D8DEE9", 2)
        rounded(d, (x+24, 315, x+82, 373), 29, color)
        d.text((x+44, 327), n, font=font(24, bold=True), fill=rgb(WHITE))
        d.text((x+24, 410), title, font=font(31, display=True), fill=rgb(NAVY))
        text_wrap(d, desc, (x+24, 465), 280, font(23), MUTED, 9)
        rounded(d, (x+24, 590, x+230, 642), 18, color)
        d.text((x+43, 604), status, font=font(18, bold=True), fill=rgb(WHITE))
    for x in (440, 845, 1250):
        arrow(d, (x, 485), (x+66, 485), DEEP, 5, 16)
    d.line((1490, 720, 1490, 835, 280, 835, 280, 705), fill=rgb(ORANGE), width=8)
    arrow(d, (280, 705), (280, 675), ORANGE, 8, 22)
    d.text((590, 785), "O ciclo ainda não volta sozinho", font=font(33, display=True), fill=rgb(ORANGE))
    text_wrap(d, "É por isso que lançar funciona, mas operar ainda parece incompleto.", (590, 840), 700, font(25), MUTED, 8)
    return save(im, "03_sistema_hoje.png")


def draw_traffic_cycle():
    im, d = canvas("A Camada de Tráfego é um ciclo", "O produto não termina quando a campanha nasce.")
    labels = ["PLANEJAR", "MONTAR", "PROVAR", "LANÇAR", "OBSERVAR", "DIAGNOSTICAR", "DECIDIR", "APRENDER"]
    colors = [GREEN, GREEN, GREEN, GREEN, ORANGE, ORANGE, PURPLE, PURPLE]
    cx, cy, r = 900, 560, 355
    pts = []
    for i in range(8):
        a = -math.pi/2 + i*2*math.pi/8
        pts.append((cx+r*math.cos(a), cy+r*math.sin(a)))
    for i, p in enumerate(pts):
        arrow(d, p, pts[(i+1)%8], "AAB5C5", 6, 18)
    for i, ((x, y), label, color) in enumerate(zip(pts, labels, colors)):
        rounded(d, (x-118, y-42, x+118, y+42), 22, WHITE, color, 4)
        tw = d.textbbox((0,0), label, font=font(20, bold=True))[2]
        d.text((x-tw/2, y-13), label, font=font(20, bold=True), fill=rgb(color))
    rounded(d, (650, 425, 1150, 685), 45, NAVY)
    d.text((745, 470), "SEARCH V1", font=font(42, display=True), fill=rgb(WHITE))
    d.text((720, 535), "forte no nascimento", font=font(27), fill=rgb(CYAN))
    d.text((718, 585), "fraca no pós-lançamento", font=font(27), fill=rgb(ORANGE))
    d.text((625, 955), "Verde = capacidade atual forte   •   Laranja/Roxo = próximo fechamento", font=font(21), fill=rgb(MUTED))
    return save(im, "04_ciclo_trafego.png")


def draw_clusters():
    im, d = canvas("Seis clusters para organizar sem perder nada", "Os clusters agrupam intenção. As ondas definem a ordem.")
    items = [
        ("A", "TERMINAR SEARCH V1", "Nascimento + retorno + histórico", CYAN),
        ("B", "OPERAR O QUE SUBIU", "Saúde, alertas e portfólio", ORANGE),
        ("C", "VERDADE DOS DADOS", "Custo, receita e atribuição", PURPLE),
        ("D", "DECISÃO GOVERNADA", "Propor, autorizar, executar e provar", DEEP),
        ("E", "INTELIGÊNCIA", "Aprender, simular e prever", CYAN),
        ("F", "EXPANSÃO", "Contrato comum e novos canais", ORANGE),
    ]
    positions = [(100,240),(650,240),(1200,240),(100,590),(650,590),(1200,590)]
    for (x,y),(letter,title,desc,color) in zip(positions, items):
        rounded(d,(x,y,x+500,y+270),30,WHITE,"D8DEE9",2)
        rounded(d,(x+28,y+28,x+100,y+100),36,color)
        d.text((x+52,y+42),letter,font=font(28,bold=True),fill=rgb(WHITE))
        text_wrap(d,title,(x+28,y+126),440,font(28,display=True),NAVY,7)
        text_wrap(d,desc,(x+28,y+190),440,font(22),MUTED,7)
    return save(im, "05_clusters.png")


def draw_matrix():
    im, d = canvas("Matriz de prioridade", "Julgamento qualitativo para ordenar trabalho — não é medição de performance.")
    x0,y0,x1,y1 = 260,230,1580,870
    d.rectangle((x0,y0,x1,y1),fill=rgb(WHITE),outline=rgb("CBD5E1"),width=3)
    d.line((920,y0,920,y1),fill=rgb("CBD5E1"),width=3)
    d.line((x0,550,x1,550),fill=rgb("CBD5E1"),width=3)
    d.text((300,255),"PREPARAR",font=font(22,bold=True),fill=rgb(MUTED))
    d.text((965,255),"FAZER AGORA",font=font(22,bold=True),fill=rgb(GREEN))
    d.text((300,585),"ESTACIONAR",font=font(22,bold=True),fill=rgb(MUTED))
    d.text((965,585),"GANHOS RÁPIDOS",font=font(22,bold=True),fill=rgb(ORANGE))
    bubbles = [
        (1180,365,"Search V1\ncompleto",GREEN,95),(1390,440,"Saúde +\nalertas",GREEN,82),
        (1010,455,"Detalhe da\ncampanha",GREEN,76),(760,390,"Custo +\nreceita",PURPLE,88),
        (520,465,"Propostas +\nautorização",PURPLE,84),(610,690,"Forecasting",DEEP,72),
        (430,760,"Rede massiva\nde personas",DEEP,78),(1160,705,"Polimento\nvisual",ORANGE,64),
        (1390,750,"Experimento\nde copy",ORANGE,68),(760,735,"Novos\ncanais",DEEP,75),
    ]
    for x,y,label,color,r in bubbles:
        d.ellipse((x-r,y-r,x+r,y+r),fill=rgb(color))
        lines=label.split("\n")
        yy=y-22 if len(lines)==2 else y-10
        for line in lines:
            tw=d.textbbox((0,0),line,font=font(19,bold=True))[2]
            d.text((x-tw/2,yy),line,font=font(19,bold=True),fill=rgb(WHITE)); yy+=27
    arrow(d,(225,870),(225,255),DEEP,5,17)
    d.text((70,555),"IMPACTO",font=font(22,bold=True),fill=rgb(DEEP))
    arrow(d,(260,920),(1570,920),DEEP,5,17)
    d.text((790,945),"RESULTADO VISÍVEL",font=font(22,bold=True),fill=rgb(DEEP))
    return save(im, "06_matriz.png")


def draw_waves():
    im, d = canvas("Roadmap por ondas", "Sem calendário artificial: uma onda termina quando sua prova de saída existe.")
    waves = [
        ("01","FECHAR SEARCH V1","Campanha volta ao VOLC com contexto, estado e histórico.",CYAN),
        ("02","OPERAR","Saúde e portfólio mostram o que pede atenção.",ORANGE),
        ("03","MEDIR","Custo, receita e atribuição fecham a verdade.",PURPLE),
        ("04","GOVERNAR","Mudanças viram propostas autorizadas e auditáveis.",DEEP),
        ("05","ESCALAR","Novos canais entram sobre um contrato comum.",GREEN),
        ("06","ANTECIPAR","Simulação e previsão entram após dados confiáveis.",CYAN),
    ]
    y=225
    for i,(n,title,desc,color) in enumerate(waves):
        x=100+i*275
        rounded(d,(x,y,x+235,850),26,WHITE,"D8DEE9",2)
        d.rectangle((x,y,x+235,y+15),fill=rgb(color))
        d.text((x+25,y+45),n,font=font(24,bold=True),fill=rgb(color))
        text_wrap(d,title,(x+25,y+105),185,font(28,display=True),NAVY,8)
        text_wrap(d,desc,(x+25,y+280),185,font(21),MUTED,8)
        d.text((x+25,790),"PROVA DE SAÍDA",font=font(16,bold=True),fill=rgb(color))
        if i<5: arrow(d,(x+235,535),(x+270,535),"AAB5C5",4,13)
    return save(im, "07_ondas.png")


def draw_bpmn():
    im, d = canvas("Fluxo-alvo: da oportunidade ao aprendizado", "O dono mantém decisões sensíveis. O sistema organiza, prova e vigia.")
    lanes=[("VOLC O.S.",220,460,PALE_BLUE),("DONO",460,655,PALE_ORANGE),("GOOGLE ADS",655,850,PALE_PURPLE)]
    for name,y1,y2,fill in lanes:
        d.rectangle((90,y1,1710,y2),fill=rgb(fill),outline=rgb("CBD5E1"),width=2)
        d.text((110,y1+20),name,font=font(18,bold=True),fill=rgb(DEEP))
    steps=[
        (220,340,"Seleciona\noportunidade",DEEP),(455,340,"Monta +\nprova",DEEP),
        (690,555,"Autoriza\nlançamento",ORANGE),(925,750,"Cria pausada\n+ revisa",PURPLE),
        (1160,555,"Decide\nativação",ORANGE),(1395,340,"Vigia +\nmede",GREEN),
        (1590,340,"Propõe\npróxima ação",DEEP),
    ]
    for x,y,label,color in steps:
        rounded(d,(x-85,y-55,x+85,y+55),24,WHITE,color,4)
        yy=y-28
        for line in label.split("\n"):
            tw=d.textbbox((0,0),line,font=font(19,bold=True))[2]
            d.text((x-tw/2,yy),line,font=font(19,bold=True),fill=rgb(color)); yy+=29
    for (x1,y1,_l1,_c1),(x2,y2,_l2,_c2) in zip(steps,steps[1:]):
        arrow(d,(x1+86,y1),(x2-86,y2),"70819A",5,15)
    d.polygon([(1520,930),(1570,880),(1620,930),(1570,980)],fill=rgb(WHITE),outline=rgb(ORANGE))
    d.text((1380,945),"Mudança relevante? volta para autorização",font=font(20,bold=True),fill=rgb(ORANGE))
    arrow(d,(1590,395),(1570,875),ORANGE,4,14)
    return save(im, "08_bpmn.png")


def draw_dependencies():
    im, d = canvas("Grafo de correspondências", "O avançado depende do básico bem fechado — não foi descartado, apenas ganhou pré-requisitos.", dark=True)
    nodes=[
        (180,520,"SEARCH V1\nCOMPLETO",CYAN),(500,300,"SAÚDE +\nALERTAS",ORANGE),
        (500,740,"CUSTO +\nRECEITA",PURPLE),(850,300,"PORTFÓLIO",CYAN),
        (850,740,"ATRIBUIÇÃO",PURPLE),(1190,520,"PROPOSTAS +\nEXECUÇÃO",ORANGE),
        (1530,300,"MULTICANAL",GREEN),(1530,740,"ORAKUL +\nFORECASTING",DEEP),
    ]
    edges=[(0,1),(0,2),(1,3),(2,3),(2,4),(3,5),(4,5),(5,6),(5,7),(4,7)]
    for a,b in edges:
        arrow(d,(nodes[a][0]+85,nodes[a][1]),(nodes[b][0]-85,nodes[b][1]),"526581",4,14)
    for x,y,label,color in nodes:
        rounded(d,(x-105,y-62,x+105,y+62),24,"151E35",color,3)
        yy=y-30
        for line in label.split("\n"):
            tw=d.textbbox((0,0),line,font=font(19,bold=True))[2]
            d.text((x-tw/2,yy),line,font=font(19,bold=True),fill=rgb(WHITE)); yy+=29
    return save(im, "09_dependencias.png")


def draw_parking():
    im, d = canvas("Estacionamento consciente", "Ideias importantes que não devem competir com o fechamento do ciclo atual.")
    items=[
        ("PREVISÃO / CRYSTAL BALL","Depende de histórico limpo e reconciliado."),
        ("ORAKUL AUTÔNOMO","Depende de propostas, autorização, replay e prova."),
        ("PMAX / META / OUTROS","Depende de um contrato multicanal e operação Search estável."),
        ("REDE MASSIVA DE PERSONAS","Depende de uma tese operacional validada, não só arquitetura."),
        ("PRICING GAM AVANÇADO","Depende da maturidade de monetização e dados de inventário."),
        ("AUTOMAÇÕES LEGADAS N8N","Precisam de dono, credenciais, observabilidade e finalidade atual."),
    ]
    for i,(title,desc) in enumerate(items):
        col=i%2; row=i//2; x=110+col*820; y=230+row*245
        rounded(d,(x,y,x+750,y+190),28,WHITE,"D8DEE9",2)
        rounded(d,(x+25,y+28,x+78,y+81),26,NAVY)
        d.text((x+44,y+39),"P",font=font(19,bold=True),fill=rgb(CYAN))
        d.text((x+105,y+30),title,font=font(24,display=True),fill=rgb(NAVY))
        text_wrap(d,desc,(x+105,y+85),600,font(21),MUTED,8)
    return save(im, "10_estacionamento.png")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in edges:
            tag = "w:" + edge
            el = borders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                borders.append(el)
            for key, value in edges[edge].items():
                el.set(qn("w:" + key), str(value))


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("VOLC O.S.  •  ")
    run.font.name = "Inter"; run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(MUTED)
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def style_doc(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Inter"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12
    for name,size,color in [("Title",34,NAVY),("Heading 1",24,NAVY),("Heading 2",16,DEEP),("Heading 3",12,ORANGE)]:
        st=styles[name]; st.font.name="Space Grotesk"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(7); st.paragraph_format.keep_with_next=True
    styles["Subtitle"].font.name="Inter"; styles["Subtitle"].font.size=Pt(13); styles["Subtitle"].font.color.rgb=RGBColor.from_string(MUTED)
    for section in doc.sections:
        section.top_margin=Cm(1.7); section.bottom_margin=Cm(1.6); section.left_margin=Cm(1.8); section.right_margin=Cm(1.8)


def configure_section(section, cover=False):
    section.page_width=Cm(21); section.page_height=Cm(29.7)
    if cover:
        section.top_margin=section.bottom_margin=section.left_margin=section.right_margin=Cm(0)
        section.header_distance=section.footer_distance=Cm(0)
    else:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.top_margin=Cm(1.6); section.bottom_margin=Cm(1.45); section.left_margin=Cm(1.75); section.right_margin=Cm(1.75)
        section.header_distance=Cm(0.7); section.footer_distance=Cm(0.65)
        h=section.header.paragraphs[0]
        h.text="WORKBOOK • PRÓXIMOS PASSOS"
        h.runs[0].font.name="Space Grotesk"; h.runs[0].font.size=Pt(8); h.runs[0].font.bold=True; h.runs[0].font.color.rgb=RGBColor.from_string(DEEP)
        set_page_number(section.footer.paragraphs[0])


def add_title(doc, title, subtitle=None, eyebrow=None):
    if eyebrow:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(eyebrow.upper()); r.bold=True; r.font.name="Inter"; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(ORANGE)
    doc.add_heading(title, level=1)
    if subtitle:
        p=doc.add_paragraph(subtitle, style="Subtitle"); p.paragraph_format.space_after=Pt(11)


def add_image(doc, path, width=17.5):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(5)
    p.add_run().add_picture(str(path), width=Cm(width))


def add_caption(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)


def callout(doc, title, text, color=DEEP, fill=PALE_BLUE):
    table=doc.add_table(rows=1, cols=2); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    table.columns[0].width=Cm(.25); table.columns[1].width=Cm(16.8)
    set_cell_shading(table.cell(0,0), color); set_cell_shading(table.cell(0,1), fill)
    for c in table.rows[0].cells: set_cell_margins(c,160,170,160,170); set_cell_border(c,top={"val":"nil"},bottom={"val":"nil"},left={"val":"nil"},right={"val":"nil"})
    p=table.cell(0,1).paragraphs[0]
    r=p.add_run(title+"\n"); r.bold=True; r.font.name="Space Grotesk"; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(color)
    r=p.add_run(text); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def bullet(doc, text, level=0, checkbox=False):
    p=doc.add_paragraph(style="List Bullet" if not checkbox else None)
    p.paragraph_format.left_indent=Cm(.55+level*.4); p.paragraph_format.first_line_indent=Cm(-.25); p.paragraph_format.space_after=Pt(3)
    if checkbox: p.add_run("☐  ").font.color.rgb=RGBColor.from_string(DEEP)
    p.add_run(text)
    return p


def simple_table(doc, headers, rows, widths=None, header_color=DEEP, font_size=8.7):
    table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    if widths:
        for i,w in enumerate(widths): table.columns[i].width=Cm(w)
    hdr=table.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        c=hdr.cells[i]; set_cell_shading(c,header_color); set_cell_margins(c)
        p=c.paragraphs[0]; r=p.add_run(h); r.bold=True; r.font.name="Inter"; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; set_cell_shading(c,WHITE if ri%2==0 else "F8FAFC"); set_cell_margins(c)
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=c.paragraphs[0]; r=p.add_run(str(val)); r.font.name="Inter"; r.font.size=Pt(font_size); r.font.color.rgb=RGBColor.from_string(INK)
            set_cell_border(c,bottom={"val":"single","sz":"4","color":"E2E8F0"})
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return table


def page_break(doc):
    doc.add_page_break()


def add_section_divider(doc, number, title, statement, color):
    table=doc.add_table(rows=1,cols=1); table.autofit=False; table.columns[0].width=Cm(17.3)
    c=table.cell(0,0); set_cell_shading(c,NAVY); set_cell_margins(c,850,650,850,650)
    p=c.paragraphs[0]
    r=p.add_run(f"{number}\n"); r.font.name="Space Grotesk"; r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(color)
    r=p.add_run(title+"\n"); r.font.name="Space Grotesk"; r.font.size=Pt(31); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE)
    r=p.add_run(statement); r.font.name="Inter"; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string("C8D2E3")
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def build_doc(images):
    doc=Document(); style_doc(doc)
    sec=doc.sections[0]; configure_section(sec,cover=True)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(images["cover"]), width=Cm(21), height=Cm(29.7))

    sec=doc.add_section(WD_SECTION.NEW_PAGE); configure_section(sec)
    add_title(doc,"Como usar este workbook","Leia primeiro para decidir; só depois converta as ondas escolhidas em tarefas.","Orientação")
    callout(doc,"A mensagem central","Você não precisa diminuir a ambição. Precisa separar horizonte, capacidade e próxima prova. O sistema já tem uma base forte; o ciclo de Tráfego é que ainda para cedo demais.",ORANGE,PALE_ORANGE)
    simple_table(doc,["LEITURA","O QUE RESPONDE","O QUE FAZER"],[
        ("Parte 1 — Norte","Que empresa e sistema estamos construindo?","Alinhar linguagem e limites."),
        ("Parte 2 — Presente","O que existe e por que ainda parece incompleto?","Aceitar a verdade atual sem desvalorizar o avanço."),
        ("Parte 3 — Ordem","O que vem agora e do que depende?","Escolher a próxima onda, não o universo inteiro."),
        ("Parte 4 — Execução","Como saber que uma etapa terminou?","Usar provas de saída e decisões explícitas."),
    ],[3.2,7.2,6.5])
    doc.add_heading("Legenda de confiança",level=2)
    simple_table(doc,["MARCA","SIGNIFICADO"],[
        ("PROVADO","Existe no sistema e há evidência de funcionamento."),
        ("CONSTRUÍDO","Existe em código, mas ainda precisa de fechamento operacional ou prova completa."),
        ("DEFINIDO","Está descrito nas fontes, ainda não é capacidade atual."),
        ("HISTÓRICO","É memória útil; precisa ser revalidada antes de virar backlog."),
        ("DECISÃO DO DONO","Depende de uma escolha de negócio, risco ou autorização."),
    ],[3.4,13.5])
    callout(doc,"Regra de leitura","Números de testes neste documento são medições do repositório em 22/08/2026. Pontuações de prioridade são julgamento qualitativo, não resultado de mídia.",DEEP,PALE_BLUE)

    page_break(doc); add_section_divider(doc,"01","Norte comum","O VOLC O.S. é a camada operacional de uma máquina de atenção, ativos, monetização, dados e audiência própria.",CYAN)
    page_break(doc); add_title(doc,"Uma visão, cinco fontes","A hierarquia evita que documentos antigos virem obrigações atuais.","Fontes oficiais")
    add_image(doc,images["sources"]); add_caption(doc,"Figura 1 — Papel de cada fonte na construção deste workbook.")
    doc.add_heading("Critério adotado",level=2)
    bullet(doc,"O Publisher Global Blueprint define o norte empresarial.")
    bullet(doc,"O Plano oficial Foco Genial define a intenção operacional inicial e os bloqueios de ativação.")
    bullet(doc,"PRD e SPEC de Arbitragem definem uma hipótese sólida de produto, mas não substituem o estado real.")
    bullet(doc,"O Second Brain preserva modelos valiosos — NEXUS, ORAKUL, WCV, Crystal Ball e outros — sem transformá-los automaticamente em prioridade.")
    bullet(doc,"O sistema em execução é a verdade de hoje. O ClickUp entra na próxima revisão incremental.")

    page_break(doc); add_title(doc,"A máquina por trás do sistema","Dois jogos convivem: caixa e aprendizado no curto prazo; audiência própria e valor durável no longo prazo.","Visão VOLC")
    add_image(doc,images["global"]); add_caption(doc,"Figura 2 — Ciclo empresarial sintetizado a partir do Publisher Global Blueprint.")
    callout(doc,"Onde entra a Camada de Tráfego","Ela transforma intenção em aquisição controlada e traz a resposta do mercado. Seu trabalho só termina quando custo, receita, qualidade e aprendizado voltam para o próximo ciclo.",PURPLE,PALE_PURPLE)
    doc.add_heading("O que não muda",level=2)
    bullet(doc,"Jogo 1 — Arbitragem: caixa, validação e aprendizagem com risco limitado.")
    bullet(doc,"Jogo 2 — Publisher: ativos, relacionamento, distribuição própria e LTV.")
    bullet(doc,"VOLC O.S.: coordena os dois jogos com clareza, prova e governança.")

    page_break(doc); add_title(doc,"Mapa da operação","Dez subsistemas existem no norte; o roadmap atual escolhe apenas as peças que fecham o próximo ciclo.","Contexto completo")
    simple_table(doc,["SUBSISTEMA","FUNÇÃO EM LINGUAGEM SIMPLES","RELAÇÃO COM O CICLO ATUAL"],[
        ("Aquisição paga","Compra atenção com intenção e limites.","Foco imediato."),
        ("Monetização","Transforma sessões em receita sustentável.","Fonte de verdade do retorno."),
        ("Conteúdo","Cria ativos que merecem a visita.","Origem do funil e da experiência."),
        ("Dados & BI","Mostra o que aconteceu e por quê.","Próximo fechamento."),
        ("Infraestrutura","Mantém a operação disponível e observável.","Espinha transversal."),
        ("Compliance","Protege usuário, contas e empresa.","Portão obrigatório."),
        ("Financeiro & crédito","Define capacidade e risco de caixa.","Entra no portfólio."),
        ("IA & agentes","Acelera análise e proposta.","Depois da verdade dos dados."),
        ("Orgânico","Reduz dependência de mídia paga.","Horizonte paralelo, não bloqueador."),
        ("Produtos & moat","Converte audiência em valor próprio.","Horizonte estratégico."),
    ],[3.0,7.2,6.7],font_size=8.2)
    callout(doc,"Princípio de corte","Uma ideia pode ser importante e ainda assim não pertencer à onda atual.",ORANGE,PALE_ORANGE)

    page_break(doc); add_section_divider(doc,"02","Verdade presente","Search V1 já nasce com bons portões. O vazio está em voltar, observar, medir e decidir dentro do mesmo sistema.",ORANGE)
    page_break(doc); add_title(doc,"O que já existe","O sistema não está no zero. Há uma cadeia real do Pautador ao Google Ads.","Estado atual")
    add_image(doc,images["now"]); add_caption(doc,"Figura 3 — Síntese do sistema observado em 22/08/2026.")
    simple_table(doc,["CAPACIDADE","ESTADO","EVIDÊNCIA / LIMITE"],[
        ("Quadro de oportunidades","PROVADO","A Camada de Tráfego recebe e organiza candidatos."),
        ("Construção Search","PROVADO","Palavras-chave, lance, orçamento e anúncios editáveis."),
        ("Prova antes de subir","PROVADO","Escopo, política e consistência são verificados."),
        ("Lançamento protegido","PROVADO","Criação exige autorização e nasce pausada."),
        ("Registro da campanha","CONSTRUÍDO","Registra a campanha; contexto e URLs ainda precisam fechar como uma unidade."),
        ("Veredito de política","CONSTRUÍDO","Existe depois do lançamento, mas o retorno completo ainda não é uma jornada única."),
        ("Sino de alertas","PROVADO","Está global no topo; hoje cobre principalmente alertas de entrega recalculados."),
        ("Histórico e portfólio","FRAGMENTADO","A campanha não tem ainda uma casa completa para voltar e operar."),
        ("Custo + receita + atribuição","FRAGMENTADO","Ainda não formam uma verdade contínua por campanha e funil."),
    ],[4.3,2.7,9.9],font_size=8.1)

    page_break(doc); add_title(doc,"Por que ainda não está legal","A sensação é coerente: a experiência cobre bem o nascimento, mas ainda não cobre a vida da campanha.","Diagnóstico de produto")
    add_image(doc,images["cycle"]); add_caption(doc,"Figura 4 — Leitura qualitativa do ciclo atual; não é uma nota de performance.")
    callout(doc,"O gap em uma frase","Hoje o VOLC ajuda a campanha a nascer. A próxima versão precisa ajudá-la a voltar, explicar seu estado e orientar a próxima ação.",ORANGE,PALE_ORANGE)
    doc.add_heading("Sintomas que isso produz",level=2)
    bullet(doc,"Depois de subir, a campanha parece sair do sistema e voltar a ser um assunto do Google Ads.")
    bullet(doc,"O sino existe, mas ainda não é uma central persistente de saúde, leitura e responsabilidade.")
    bullet(doc,"Custo, receita, política, tracking e contexto não aparecem como uma única história operacional.")
    bullet(doc,"A visão avançada parece distante porque os pré-requisitos estão misturados com as ideias finais.")

    page_break(doc); add_title(doc,"Search V1 — definição de pronto","Esta é a fronteira recomendada para chamar a primeira Camada de Tráfego de completa.","Contrato do produto")
    for item in [
        "A oportunidade chega com origem, projeto, funil e responsável identificados.",
        "A conta Google Ads só pode ser escolhida dentro do escopo da casa.",
        "Palavras-chave, negativos, orçamento, lance e anúncios podem ser revisados antes do envio.",
        "A prova reúne destino publicado, tracking, política, orçamento e escopo em um único portão.",
        "O dono autoriza qualquer criação externa; a campanha nasce pausada.",
        "Campanha, grupos, anúncios, palavras-chave e URLs voltam registrados como uma única operação.",
        "O operador pode retornar a uma página da campanha e entender o que aconteceu desde o lançamento.",
        "Política, ativação, entrega, custo, conversão e receita têm estado e data de atualização visíveis.",
        "Alertas são persistentes, têm leitura/resolução e apontam para o objeto que exige atenção.",
        "Toda mudança relevante é proposta, autorizada, executada e verificada — nunca escondida.",
    ]: bullet(doc,item,checkbox=True)
    callout(doc,"Prova de saída da Onda 01","Uma campanha Search pode nascer, voltar ao VOLC no dia seguinte e ser compreendida sem reconstruir sua história em outra ferramenta.",GREEN,PALE_GREEN)

    page_break(doc); add_title(doc,"Saúde: do sino à central operacional","O sino resolve descoberta. A central de saúde resolve continuidade.","Notificações")
    simple_table(doc,["CAMADA","HOJE","PRÓXIMO FECHAMENTO"],[
        ("Presença","Sino global no topo.","Manter visível em todas as telas."),
        ("Origem","Alertas de entrega recalculados.","Somar política, tracking, dados parados e falhas de rotina."),
        ("Memória","Sem caixa persistente de lido/resolvido.","Registrar estado, primeira ocorrência, atualização e resolução."),
        ("Ação","Leitura geral.","Cada alerta abre exatamente a campanha, funil ou rotina afetada."),
        ("Responsabilidade","Implícita.","Dono, severidade, prazo operacional e escalonamento."),
        ("Proteção","Foco em mídia.","Heartbeat, rotina parada e deadman para falhas críticas."),
    ],[3.2,6.1,7.6])
    doc.add_heading("Famílias de alerta recomendadas",level=2)
    for text in ["Entrega: ativa sem impressões, gasto fora do esperado, orçamento ou lance limitante.","Política: revisão pendente, reprovação ou ativo afetado.","Tracking: clique sem identificação, destino inconsistente, conversões paradas.","Dados: custo ou receita desatualizados, diferença de moeda ou origem desconhecida.","Sistema: rotina atrasada, integração indisponível ou prova incompleta."]:
        bullet(doc,text)
    callout(doc,"Ordem correta","Persistência e contexto antes de multiplicar regras. Cem alertas sem memória criam ruído, não segurança.",DEEP,PALE_BLUE)

    page_break(doc); add_section_divider(doc,"03","Ordem de construção","Primeiro fechar Search, depois operar, medir, governar, escalar e antecipar.",PURPLE)
    page_break(doc); add_title(doc,"Clusters de capacidade","A visão completa cabe em seis caixas. Isso reduz ansiedade sem amputar a ambição.","Arquitetura do roadmap")
    add_image(doc,images["clusters"]); add_caption(doc,"Figura 5 — Clusters de capacidade; a ordem real está nas ondas seguintes.")
    simple_table(doc,["CLUSTER","ENTREGA PRINCIPAL","PRÉ-REQUISITO"],[
        ("A — Terminar Search V1","Jornada completa do nascimento ao retorno.","Base atual."),
        ("B — Operar o que subiu","Central de saúde e portfólio diário.","Objetos e estados persistentes."),
        ("C — Verdade dos dados","Custo, receita, moeda, origem e frescor.","Campanhas e funis reconciliados."),
        ("D — Decisão governada","Propostas, autorizações, execução e prova.","Dados confiáveis e trilha de auditoria."),
        ("E — Inteligência","Replay, shadow, aprendizado e previsão.","Histórico suficiente e decisões rotuladas."),
        ("F — Expansão","Novos canais sobre contrato comum.","Search operável e abstrações validadas."),
    ],[4.0,7.0,5.9],font_size=8.4)

    page_break(doc); add_title(doc,"Impacto versus resultado visível","A matriz ajuda a decidir o que abre o próximo ciclo e o que apenas parece avançado.","Priorização")
    add_image(doc,images["matrix"]); add_caption(doc,"Figura 6 — Avaliação editorial qualitativa baseada nas dependências das fontes e no estado observado.")
    callout(doc,"Leitura recomendada","Search V1 completo, detalhe da campanha e saúde persistente ficam no quadrante “fazer agora”. Custo e receita têm impacto alto, mas pedem preparação de dados. Forecasting e rede massiva de personas ficam estacionados.",ORANGE,PALE_ORANGE)

    page_break(doc); add_title(doc,"Roadmap por ondas","A onda seguinte não começa porque a anterior ficou cansativa; começa quando a prova de saída existe.","Sequência recomendada")
    add_image(doc,images["waves"]); add_caption(doc,"Figura 7 — Sequência sem estimativas artificiais de calendário.")
    simple_table(doc,["ONDA","ESCOPO MÍNIMO","PROVA DE SAÍDA"],[
        ("01 — Fechar Search V1","Detalhe da campanha; registro transacional de contexto e URLs; retorno pós-lançamento; estados claros.","Uma campanha pode voltar ao VOLC e ser compreendida."),
        ("02 — Operar","Central persistente; famílias de alerta; portfólio; heartbeat.","O sistema aponta o que pede atenção e por quê."),
        ("03 — Medir","Custo, receita, moeda, frescor, reconciliação e atribuição.","Cada campanha mostra resultado com fonte e data."),
        ("04 — Governar","Proposta; autorização assimétrica; executor; recibo; verificação.","Nenhuma mudança relevante acontece sem trilha."),
        ("05 — Escalar","Contrato de canal e primeiro novo adaptador escolhido.","Um segundo canal reutiliza o núcleo sem copiar a lógica."),
        ("06 — Antecipar","Replay, shadow, aprendizado e forecasting com confiança explícita.","Uma recomendação pode ser explicada e comparada ao ocorrido."),
    ],[3.0,8.0,5.9],font_size=8.0)

    page_break(doc); add_title(doc,"Onda 01 — pacote de trabalho","O foco de curto prazo é encerrar a sensação de campanha “jogada para fora” depois do lançamento.","Backlog orientado a resultado")
    simple_table(doc,["FRENTE","RESULTADO PARA O USUÁRIO","PRÉ-REQUISITO","PRONTO QUANDO"],[
        ("Casa da campanha","Abrir uma campanha e ver sua história.","Rota e identidade estável.","Existe página de detalhe e retorno pelo quadro/alerta."),
        ("Registro único","Não perder origem, funil, URLs e configuração.","Modelo de registro definido.","Criação externa e registro local fecham com recibo claro."),
        ("Pós-lançamento","Ver política, ativação e entrega no mesmo lugar.","Leitura Google + estado local.","Estados têm fonte, data e próxima ação."),
        ("Tracking visível","Saber se o clique consegue voltar para casa.","Contrato de parâmetros.","Destino e identificação são provados antes/depois."),
        ("Alertas com contexto","Ir do sino ao objeto correto.","Persistência mínima.","Ler, resolver e reabrir funciona."),
        ("Retomada","Continuar o trabalho sem reconstruir escolhas.","Estado de tela persistido.","Operador retorna sem perda relevante."),
    ],[2.8,4.7,4.0,5.4],font_size=7.7)
    callout(doc,"Não entra nesta onda","Otimização autônoma, previsão, rede de personas, expansão de canais e pricing avançado. Todos permanecem registrados no estacionamento.",PURPLE,PALE_PURPLE)

    page_break(doc); add_title(doc,"Fluxo-alvo de negócio","Uma versão simples do BPMN: sistema executa e prova; o dono autoriza decisões sensíveis.","BPMN simplificado")
    add_image(doc,images["bpmn"]); add_caption(doc,"Figura 8 — Fluxo-alvo; swimlanes representam responsabilidades, não componentes técnicos.")
    doc.add_heading("Portões que nunca somem",level=2)
    bullet(doc,"Destino publicado, acessível e coerente com o anúncio.")
    bullet(doc,"Consentimento, políticas, ads.txt e qualidade de página quando aplicáveis ao ativo.")
    bullet(doc,"Conta, orçamento, moeda, escopo e autorização explícitos.")
    bullet(doc,"Campanha nasce pausada; ativação e aumentos relevantes pertencem ao dono.")
    bullet(doc,"Toda execução retorna recibo e toda leitura informa quando foi atualizada.")

    page_break(doc); add_title(doc,"Dependências: por que a ordem importa","As ideias avançadas continuam válidas, mas agora sabemos qual chão cada uma exige.","Grafo de correspondências")
    add_image(doc,images["deps"]); add_caption(doc,"Figura 9 — Grafo lógico de pré-requisitos, sintetizado das fontes.")
    callout(doc,"Correspondência importante","ORAKUL deixa de ser “um agente mágico” e passa a ser a camada que compara fatos, simulações e regras para produzir uma proposta. Sem dados, autorização e recibo, ele não deve executar.",PURPLE,PALE_PURPLE)
    simple_table(doc,["VISÃO HISTÓRICA","CAPACIDADE MODERNA CORRESPONDENTE","QUANDO REABRIR"],[
        ("NEXUS","Intenção da sessão + próximo melhor ativo.","Após atribuição e ativos instrumentados."),
        ("ORAKUL","Motor de proposta explicável.","Após dados e governança."),
        ("Crystal Ball","Forecast com faixa de confiança.","Após histórico reconciliado."),
        ("WCV / valor ponderado","Sensor de qualidade de tráfego e monetização.","Após conversões confiáveis."),
        ("BEAST","Economia CPC × valor da sessão.","Após custo e receita na mesma granularidade."),
        ("Workflows n8n","Rotinas operacionais observáveis.","Após dono, credencial e finalidade validados."),
    ],[4.2,7.0,5.7],font_size=8.0)

    page_break(doc); add_title(doc,"Estacionamento consciente","Guardar com critério é diferente de esquecer.","Não abrir agora")
    add_image(doc,images["parking"]); add_caption(doc,"Figura 10 — Iniciativas preservadas fora da disputa da Onda 01.")
    callout(doc,"Regra para sair do estacionamento","A iniciativa precisa declarar problema atual, pré-requisitos disponíveis, dono, prova de saída e impacto esperado. Ambição sem esses cinco campos continua como referência.",DEEP,PALE_BLUE)

    page_break(doc); add_section_divider(doc,"04","Execução com calma","Converter uma onda em decisões, provas e poucas frentes abertas.",GREEN)
    page_break(doc); add_title(doc,"Quadro de decisões do dono","As escolhas abaixo mudam risco, escopo ou negócio; o sistema não deve inventá-las.","Decisões abertas")
    decisions=[
        ("Fronteira da Search V1","A definição de pronto deste workbook está aprovada?","☐ Sim  ☐ Ajustar"),
        ("Ativação de campanhas","Quem pode ativar e em quais condições?","________________________________"),
        ("Publicação de páginas","Qual é o portão mínimo de conteúdo/compliance?","________________________________"),
        ("Conversões","Qual evento será fonte oficial no primeiro ciclo?","________________________________"),
        ("Economia","Qual moeda e qual visão de receita entram primeiro?","________________________________"),
        ("Alertas","Quem recebe, resolve e escala cada severidade?","________________________________"),
        ("Próximo canal","Qual dor real justifica sair de Search?","________________________________"),
        ("Legado","Quais rotinas antigas continuam autorizadas?","________________________________"),
    ]
    simple_table(doc,["DECISÃO","PERGUNTA","RESPOSTA"],decisions,[4.0,8.0,4.9],font_size=8.2)
    callout(doc,"Sugestão de governança","Uma decisão aberta pode bloquear uma tarefa. Uma tarefa nunca deve esconder uma decisão aberta.",ORANGE,PALE_ORANGE)

    page_break(doc); add_title(doc,"Ficha de iniciativa","Use uma página como esta para cada frente que entrar na onda ativa.","Template de execução")
    fields=[
        ("NOME DA INICIATIVA","____________________________________________________________"),
        ("PROBLEMA QUE RESOLVE","____________________________________________________________\n____________________________________________________________"),
        ("RESULTADO VISÍVEL PARA O USUÁRIO","____________________________________________________________\n____________________________________________________________"),
        ("PRÉ-REQUISITOS","☐ Dados  ☐ Decisão  ☐ Integração  ☐ Design  ☐ Compliance"),
        ("FORA DO ESCOPO","____________________________________________________________"),
        ("PROVA DE SAÍDA","____________________________________________________________\n____________________________________________________________"),
        ("FONTE / CORRESPONDÊNCIA","____________________________________________________________"),
        ("DONO E PRÓXIMA REVISÃO","____________________________________________________________"),
    ]
    t=doc.add_table(rows=0,cols=1); t.autofit=False; t.columns[0].width=Cm(17.2)
    for label,val in fields:
        c=t.add_row().cells[0]; set_cell_shading(c,WHITE); set_cell_margins(c,145,180,170,180); set_cell_border(c,bottom={"val":"single","sz":"8","color":"DDE4EE"})
        p=c.paragraphs[0]; r=p.add_run(label+"\n"); r.bold=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(DEEP)
        r=p.add_run(val); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(MUTED)

    page_break(doc); add_title(doc,"Ritual de foco","Uma cadência simples para manter a visão sem voltar à bagunça.","Modo de operação")
    simple_table(doc,["MOMENTO","PERGUNTAS","SAÍDA"],[
        ("Abrir a onda","Qual prova precisa existir? Quais decisões estão abertas?","Uma fronteira de pronto."),
        ("Escolher frentes","O que pode avançar sem abrir dependências invisíveis?","Poucas iniciativas ativas."),
        ("Revisar evidência","O que foi observado? Qual fonte e data?","Fato separado de opinião."),
        ("Encerrar","A prova de saída existe? O usuário percebe o resultado?","Onda concluída ou gap explícito."),
        ("Promover próxima","Os pré-requisitos da onda seguinte existem?","Nova onda ativa."),
    ],[3.0,9.0,4.9])
    doc.add_heading("Regra dos três horizontes",level=2)
    bullet(doc,"ATIVO — uma onda, com poucas frentes e prova de saída clara.")
    bullet(doc,"PRÓXIMO — uma onda preparada, ainda sem competir por execução.")
    bullet(doc,"ESTACIONADO — todo o resto, preservado com correspondência e pré-requisitos.")
    callout(doc,"Para puxar o ar","A complexidade deixa de morar na sua cabeça quando cada ideia sabe em qual horizonte está e o que precisa acontecer antes dela.",CYAN,PALE_BLUE)

    page_break(doc); add_title(doc,"Plano de partida","O primeiro movimento recomendado após aprovar este workbook.","Próxima sessão")
    for i,(title,desc) in enumerate([
        ("Congelar a fronteira da Search V1","Revisar a checklist de pronto e ajustar apenas o que muda a experiência do operador."),
        ("Abrir uma iniciativa por frente da Onda 01","Casa da campanha, registro único, pós-lançamento, tracking visível, alertas com contexto e retomada."),
        ("Transformar decisões abertas em respostas","Especialmente ativação, conversão oficial, economia e responsabilidade dos alertas."),
        ("Levar somente a Onda 01 ao ClickUp","Criar tarefas executáveis; manter ondas futuras como épicos ou referências, não como fila concorrente."),
        ("Revisar após a primeira campanha que completar o ciclo","A prova real decide a Onda 02."),
    ],1):
        table=doc.add_table(rows=1,cols=2); table.autofit=False; table.columns[0].width=Cm(1.3); table.columns[1].width=Cm(15.8)
        set_cell_shading(table.cell(0,0),DEEP); set_cell_shading(table.cell(0,1),WHITE)
        for c in table.rows[0].cells: set_cell_margins(c,150,170,150,170); set_cell_border(c,bottom={"val":"single","sz":"7","color":"DDE4EE"})
        p=table.cell(0,0).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(i)); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(WHITE)
        p=table.cell(0,1).paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.name="Space Grotesk"; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(NAVY)
        r=p.add_run(desc); r.font.size=Pt(9.5); r.font.color.rgb=RGBColor.from_string(MUTED)
    callout(doc,"Resultado da próxima sessão","Sair com a Search V1 fechada como produto e a Onda 01 pronta para execução — sem tentar implementar o roadmap inteiro.",GREEN,PALE_GREEN)

    page_break(doc); add_title(doc,"Medições e validações desta edição","Registro objetivo do que foi verificado no repositório em 22/08/2026.","Evidências")
    simple_table(doc,["VERIFICAÇÃO","RESULTADO MEDIDO","LEITURA"],[
        ("Testes da interface de Tráfego, sino e hooks","93 testes aprovados em 13 arquivos.","A base de interface consultada passou no conjunto direcionado."),
        ("Build da interface","Build concluído; 4.303 módulos transformados.","O pacote de frontend compilou no ambiente local."),
        ("Testes backend direcionados","63 aprovados, 6 ignorados, 4 falhas e 5 erros.","Falhas/erros ocorreram em tentativas externas bloqueadas por rede no ambiente; não foram tratados como defeito funcional comprovado."),
        ("Prova visual do sino","Sino global observado no topo da interface.","Capacidade atual confirmada; central persistente ainda é gap."),
    ],[5.0,5.4,6.5],font_size=8.0)
    callout(doc,"Limite desta edição","Não houve nova leitura ao vivo das contas de mídia em 22/08/2026. Evidências de campanhas observadas em 20/08/2026 permanecem históricas e não são apresentadas como estado atual.",ORANGE,PALE_ORANGE)

    page_break(doc); add_title(doc,"Fontes e rastreabilidade","Documento-base, referências estratégicas e verdade atual consultados nesta edição.","Apêndice")
    sources=[
        ("Plano oficial de otimização performance e monetização — Foco Genial.docx","Fonte inicial oficial; versão 1.1, 23/07/2026."),
        ("docs/PRD-ARBITRAGEM.md","Sequência de evolução do motor de arbitragem."),
        ("docs/SPEC-ARBITRAGEM.md","Princípios de saúde, dados, propostas, autorização e execução."),
        ("Publisher Global Blueprint","Norte empresarial, dez subsistemas e fases globais."),
        ("SCOND-BRAIN-MEDIA-ARBITRAGE","Memória de modelos e experimentos: VOS PRIME, NEXUS, ORAKUL, Crystal Ball, BEAST, WCV e workflows."),
        ("VOLC O.S. — repositório local","Rotas, componentes, motor Google Ads, alertas, testes e build observados em 22/08/2026."),
        ("ClickUp — lista 901328196164","Fonte incremental indicada pelo dono; conteúdo não incorporado nesta versão."),
    ]
    simple_table(doc,["FONTE","USO NESTE WORKBOOK"],sources,[8.0,8.9],font_size=8.3)
    doc.add_heading("Notas editoriais",level=2)
    bullet(doc,"Termos técnicos foram traduzidos para decisões, capacidades, provas e responsabilidades.")
    bullet(doc,"Ideias históricas não foram descartadas; foram mapeadas a capacidades modernas e pré-requisitos.")
    bullet(doc,"Recomendações de prioridade são julgamento editorial baseado em dependências, não promessa de retorno.")
    bullet(doc,"Esta versão pode receber ClickUp, novas referências e decisões sem refazer sua estrutura central.")
    callout(doc,"Próxima revisão sugerida","Depois de incorporar o ClickUp e fechar as decisões abertas da Search V1, publicar a versão 0.2 com backlog aprovado e correspondência tarefa → onda → prova.",DEEP,PALE_BLUE)

    page_break(doc)
    table=doc.add_table(rows=1,cols=1); table.autofit=False; table.columns[0].width=Cm(17.3)
    c=table.cell(0,0); set_cell_shading(c,NAVY); set_cell_margins(c,900,650,900,650)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("A visão continua enorme.\n"); r.font.name="Space Grotesk"; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(WHITE)
    r=p.add_run("Agora, o próximo passo cabe na mão.\n\n"); r.font.name="Space Grotesk"; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(CYAN)
    r=p.add_run("VOLC O.S.  •  Workbook de Próximos Passos"); r.font.name="Inter"; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string("B7C3D9")

    # Metadados.
    doc.core_properties.title="VOLC O.S. — Workbook de Próximos Passos"
    doc.core_properties.subject="Camada de Tráfego e roadmap do sistema"
    doc.core_properties.author="VOLC"
    doc.core_properties.keywords="VOLC O.S., tráfego, Google Ads, roadmap, workbook"
    doc.core_properties.comments="Gerado a partir das fontes oficiais e do estado observado do sistema em 22/08/2026."
    doc.save(OUT_PATH)


def main():
    ASSET_DIR.mkdir(parents=True,exist_ok=True); OUT_DIR.mkdir(parents=True,exist_ok=True)
    images={
        "cover":draw_cover(),"sources":draw_source_map(),"global":draw_global_loop(),
        "now":draw_system_now(),"cycle":draw_traffic_cycle(),"clusters":draw_clusters(),
        "matrix":draw_matrix(),"waves":draw_waves(),"bpmn":draw_bpmn(),
        "deps":draw_dependencies(),"parking":draw_parking(),
    }
    build_doc(images)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
