#!/usr/bin/env python3
"""Gera o Livro Vivo do VOLC O.S. a partir do grafo, narrativa e roadmap.

Fontes:
  - LIVRO-FONTE.md: explicação humana
  - ROADMAP-VIVO.json: iniciativas e tarefas
  - docs/volc-os-graph/volc-os-graph.json: capacidades e estados

Saída:
  - entregaveis/Workbook_VOLC_OS_Livro_Vivo_v1.0.docx
"""

from __future__ import annotations

import json
import math
import re
from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
HERE = Path(__file__).resolve().parent
GRAPH = ROOT / "docs" / "volc-os-graph" / "volc-os-graph.json"
NARRATIVE = HERE / "LIVRO-FONTE.md"
ROADMAP = HERE / "ROADMAP-VIVO.json"
ASSETS = HERE / "assets"
OUT = ROOT / "entregaveis" / "Workbook_VOLC_OS_Livro_Vivo_v1.0.docx"

NAVY = "0B1020"
DEEP = "0D47A1"
CYAN = "00BFEA"
PURPLE = "7546E8"
ORANGE = "FF4D20"
GREEN = "16845B"
YELLOW = "D89B00"
RED = "D92D20"
INK = "1A1C1E"
MUTED = "667085"
WHITE = "FFFFFF"
OFF = "F4F6F9"
LINE = "DDE4EE"
PALE_BLUE = "EAF6FF"
PALE_PURPLE = "F3ECFF"
PALE_ORANGE = "FFF0EB"
PALE_GREEN = "EAF8F1"
PALE_YELLOW = "FFF8DF"

SPACE_FONT = Path("/Users/mac/Library/Fonts/SpaceGrotesk-VariableFont_wght.ttf")
BODY_FONT = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
BODY_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
LOGO = ROOT / "public" / "volc-logo-light.png"

W, H = 1800, 1040

STATE_LABELS = {
    "live": "Vivo agora",
    "implemented": "Implementado",
    "partial": "Parcial",
    "historical": "Histórico",
    "inactive": "Inativo",
    "empty": "Estrutura vazia",
    "decision": "Decisão aberta",
    "risk": "Risco",
    "todo": "Planejado",
}

STATE_COLORS = {
    "live": GREEN,
    "implemented": CYAN,
    "partial": ORANGE,
    "historical": PURPLE,
    "inactive": MUTED,
    "empty": MUTED,
    "decision": YELLOW,
    "risk": RED,
    "todo": DEEP,
}

TASK_COLORS = {
    "done": GREEN,
    "partial": ORANGE,
    "risk": RED,
    "todo": DEEP,
    "reserved": MUTED,
}

TASK_ICONS = {
    "done": "☑",
    "partial": "◐",
    "risk": "⚠",
    "todo": "☐",
    "reserved": "◇",
}

WAVE_COLORS = {
    "A — Clareza e fundação": CYAN,
    "B — Operação que gera caixa": ORANGE,
    "C — Escala governada": PURPLE,
    "D — Estacionamento consciente": MUTED,
}


def rgb(value: str) -> tuple[int, int, int]:
    return tuple(int(value[i:i + 2], 16) for i in (0, 2, 4))


def pil_font(size: int, *, bold: bool = False, display: bool = False):
    path = SPACE_FONT if display and SPACE_FONT.exists() else (BODY_BOLD if bold else BODY_FONT)
    return ImageFont.truetype(str(path), size=size)


def rounded(draw: ImageDraw.ImageDraw, xy, radius: int, fill: str, outline: str | None = None, width: int = 1):
    draw.rounded_rectangle(xy, radius=radius, fill=rgb(fill), outline=rgb(outline) if outline else None, width=width)


def wrap(draw: ImageDraw.ImageDraw, text: str, xy, max_width: int, font, fill: str, gap: int = 8, max_lines: int | None = None):
    words = text.split()
    lines, current = [], ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = lines[-1].rstrip(".,;:") + "…"
    x, y = xy
    line_height = font.getbbox("Ag")[3] + gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=rgb(fill))
        y += line_height
    return y


def arrow(draw: ImageDraw.ImageDraw, start, end, color: str, width: int = 6, head: int = 18):
    draw.line((start, end), fill=rgb(color), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    points = [
        end,
        (end[0] - head * math.cos(angle - .48), end[1] - head * math.sin(angle - .48)),
        (end[0] - head * math.cos(angle + .48), end[1] - head * math.sin(angle + .48)),
    ]
    draw.polygon(points, fill=rgb(color))


def save_asset(image: Image.Image, name: str) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name
    image.save(path, quality=96)
    return path


def progress_for_tasks(tasks: list[dict], weights: dict) -> tuple[float, int, int]:
    total = 0.0
    earned = 0.0
    included = 0
    for task in tasks:
        weight = weights.get(task["status"])
        if weight is None:
            continue
        included += 1
        total += 1.0
        earned += float(weight)
    return ((earned / total) if total else 0.0), included, len(tasks) - included


def roadmap_stats(roadmap: dict) -> dict:
    all_tasks = [task for initiative in roadmap["initiatives"] for task in initiative["tasks"]]
    progress, included, reserved = progress_for_tasks(all_tasks, roadmap["status_weights"])
    statuses = Counter(task["status"] for task in all_tasks)
    return {
        "progress": progress,
        "included": included,
        "reserved": reserved,
        "total": len(all_tasks),
        "statuses": statuses,
    }


def draw_cover(stats: dict, graph: dict) -> Path:
    image = Image.new("RGB", (1400, 1980), rgb(NAVY))
    glow = Image.new("RGBA", image.size, (0, 0, 0, 0))
    gd = ImageDraw.Draw(glow)
    for cx, cy, color, radius, alpha in [
        (1190, 350, CYAN, 430, 120),
        (1070, 700, PURPLE, 540, 110),
        (1260, 1040, ORANGE, 420, 100),
        (90, 1770, DEEP, 500, 125),
    ]:
        gd.ellipse((cx-radius, cy-radius, cx+radius, cy+radius), fill=(*rgb(color), alpha))
    glow = glow.filter(ImageFilter.GaussianBlur(115))
    image = Image.alpha_composite(image.convert("RGBA"), glow).convert("RGB")
    draw = ImageDraw.Draw(image)
    if LOGO.exists():
        logo = Image.open(LOGO).convert("RGBA")
        logo.thumbnail((390, 150))
        image.paste(logo, (90, 80), logo)
    draw.text((95, 430), "LIVRO VIVO · EDIÇÃO 1.0", font=pil_font(37, display=True), fill=rgb(CYAN))
    draw.text((90, 520), "VOLC O.S.", font=pil_font(112, display=True), fill=rgb(WHITE))
    draw.text((95, 690), "A operação inteira", font=pil_font(70, display=True), fill=rgb(WHITE))
    draw.text((95, 782), "em uma ordem possível", font=pil_font(60, display=True), fill=rgb(WHITE))
    wrap(draw, "Visão simples · mapa factual · tarefas prioritárias · checklists de saída",
         (100, 920), 1080, pil_font(31), WHITE, 11)

    rounded(draw, (95, 1160, 1120, 1485), 30, "141D36", CYAN, 3)
    draw.text((135, 1208), "A PROMESSA DESTA EDIÇÃO", font=pil_font(23, bold=True), fill=rgb(CYAN))
    wrap(draw, "Nada importante é apagado. Nem tudo importante precisa estar aberto agora.",
         (135, 1270), 850, pil_font(31, display=True), WHITE, 12)
    draw.text((98, 1740), "26 de agosto de 2026", font=pil_font(22), fill=rgb("B7C3D9"))
    draw.text((98, 1790), f"{stats['total']} tarefas catalogadas · {stats['reserved']} reservadas fora do percentual",
              font=pil_font(20), fill=rgb("B7C3D9"))
    draw.text((98, 1835), f"{len(graph['nodes'])} nós operacionais · {len(graph['edges'])} relações",
              font=pil_font(20), fill=rgb("B7C3D9"))
    return save_asset(image, "00_capa.png")


def canvas(title: str, subtitle: str, dark: bool = False):
    image = Image.new("RGB", (W, H), rgb(NAVY if dark else OFF))
    draw = ImageDraw.Draw(image)
    draw.text((90, 62), title, font=pil_font(49, display=True), fill=rgb(WHITE if dark else NAVY))
    wrap(draw, subtitle, (92, 128), 1540, pil_font(24), "C8D2E3" if dark else MUTED, 6)
    return image, draw


def draw_attention_map() -> Path:
    image, draw = canvas("O Mapa da Atenção", "Uma jornada só fica completa quando origem, pouso, avanço, retorno e resultado conseguem conversar.", dark=True)
    steps = [
        ("1", "ORIGEM", "Google · Meta · orgânico", CYAN),
        ("2", "POUSO", "site · vídeo · conversa", PURPLE),
        ("3", "AVANÇO", "conteúdo · interação · isca", ORANGE),
        ("4", "RETORNO", "canal consentido", GREEN),
        ("5", "RESULTADO", "receita · ativo · aprendizado", CYAN),
    ]
    y = 470
    for i, (number, title, desc, color) in enumerate(steps):
        x = 100 + i * 340
        rounded(draw, (x, y-115, x+270, y+115), 28, "151E35", color, 3)
        rounded(draw, (x+20, y-92, x+72, y-40), 26, color)
        draw.text((x+38, y-82), number, font=pil_font(20, bold=True), fill=rgb(WHITE))
        draw.text((x+22, y-12), title, font=pil_font(24, bold=True), fill=rgb(WHITE))
        wrap(draw, desc, (x+22, y+36), 225, pil_font(19), "B7C3D9", 6, 3)
        if i < len(steps) - 1:
            arrow(draw, (x+278, y), (x+330, y), "65738B", 5, 15)
    rounded(draw, (340, 790, 1460, 915), 28, "0F1930", ORANGE, 2)
    draw.text((390, 827), "Se um elo não tem fonte, data ou identidade, o mapa mostra uma lacuna — nunca inventa zero.",
              font=pil_font(25, display=True), fill=rgb(WHITE))
    return save_asset(image, "01_mapa_atencao.png")


def draw_operations() -> Path:
    image, draw = canvas("Quatro operações, uma fundação", "Separar as regras permite integrar o ecossistema sem criar um monólito.")
    center = (900, 535)
    rounded(draw, (650, 405, 1150, 665), 42, NAVY, CYAN, 4)
    draw.text((730, 455), "COFRE DE ATIVOS", font=pil_font(36, display=True), fill=rgb(WHITE))
    wrap(draw, "identidade · dono · estado · acesso · relações · próximo uso",
         (730, 525), 360, pil_font(22), "C8D2E3", 8)
    cards = [
        (90, 240, 510, 510, "MÍDIA PAGA", "Google e Meta direto para o site.", CYAN),
        (1290, 240, 1710, 510, "ORGÂNICO", "Pauta, peça, publicação e aprendizado.", PURPLE),
        (90, 650, 510, 920, "RETENÇÃO", "Permissão para a pessoa voltar.", GREEN),
        (1290, 650, 1710, 920, "NOVAS RECEITAS", "Afiliados, produtos e ferramentas.", ORANGE),
    ]
    for x1, y1, x2, y2, title, desc, color in cards:
        rounded(draw, (x1, y1, x2, y2), 30, WHITE, color, 3)
        draw.text((x1+30, y1+38), title, font=pil_font(27, bold=True), fill=rgb(color))
        wrap(draw, desc, (x1+30, y1+105), x2-x1-60, pil_font(26, display=True), NAVY, 9)
        sx = x2 if x1 < center[0] else x1
        sy = (y1+y2)//2
        ex = 640 if x1 < center[0] else 1160
        ey = 505 if y1 < center[1] else 600
        arrow(draw, (sx, sy), (ex, ey), "98A5B8", 4, 14)
    return save_asset(image, "02_operacoes.png")


def draw_asset_vault() -> Path:
    image, draw = canvas("O Cofre de Ativos", "O sistema precisa saber o que a VOLC possui antes de decidir onde publicar, anunciar ou reter.", dark=True)
    families = [
        ("SOCIAL", "páginas · perfis · grupos", CYAN),
        ("MÍDIA", "contas · pixels · apps", PURPLE),
        ("PROPRIEDADE", "domínios · sites · monetização", ORANGE),
        ("ENGINES", "texto · imagem · vídeo", GREEN),
        ("AUTOMAÇÃO", "workflows · serviços · webhooks", CYAN),
        ("ACESSO", "referência segura · owner · validade", PURPLE),
    ]
    positions = [(110,230),(660,230),(1210,230),(110,610),(660,610),(1210,610)]
    for (title, desc, color), (x, y) in zip(families, positions):
        rounded(draw, (x, y, x+480, y+260), 30, "151E35", color, 3)
        draw.text((x+30, y+38), title, font=pil_font(25, bold=True), fill=rgb(color))
        wrap(draw, desc, (x+30, y+102), 410, pil_font(26, display=True), WHITE, 9)
    return save_asset(image, "03_cofre.png")


def draw_waves(roadmap: dict) -> Path:
    image, draw = canvas("Quatro ondas para respirar e avançar", "O estacionamento preserva o futuro sem colocá-lo na sua fila diária.")
    groups = defaultdict(list)
    for item in roadmap["initiatives"]:
        groups[item["wave"]].append(item)
    order = list(WAVE_COLORS)
    for i, wave in enumerate(order):
        x = 85 + i * 430
        color = WAVE_COLORS[wave]
        rounded(draw, (x, 205, x+390, 940), 32, WHITE, color, 3)
        draw.rectangle((x, 205, x+390, 220), fill=rgb(color))
        draw.text((x+28, 252), wave.split(" — ")[0], font=pil_font(46, display=True), fill=rgb(color))
        wrap(draw, wave.split(" — ")[1], (x+28, 325), 330, pil_font(27, display=True), NAVY, 8, 3)
        y = 455
        for item in groups.get(wave, []):
            draw.text((x+28, y), f"{item['rank']:02d}", font=pil_font(18, bold=True), fill=rgb(color))
            y = wrap(draw, item["title"], (x+72, y-2), 280, pil_font(18, bold=True), INK, 5, 3) + 16
    return save_asset(image, "04_ondas.png")


def draw_progress(stats: dict) -> Path:
    image, draw = canvas("Progresso sem falsa precisão", "O índice mostra fechamento do escopo aceito — não prazo, faturamento ou dificuldade.", dark=True)
    value = stats["progress"]
    draw.arc((180, 230, 780, 830), start=135, end=405, fill=rgb("34415B"), width=52)
    draw.arc((180, 230, 780, 830), start=135, end=135 + 270*value, fill=rgb(CYAN), width=52)
    pct = f"{value*100:.0f}%"
    box = draw.textbbox((0,0), pct, font=pil_font(92, display=True))
    draw.text((480-(box[2]-box[0])/2, 455), pct, font=pil_font(92, display=True), fill=rgb(WHITE))
    draw.text((330, 580), "ÍNDICE EDITORIAL", font=pil_font(22, bold=True), fill=rgb(CYAN))
    rows = [
        ("Concluídas", stats["statuses"].get("done", 0), GREEN),
        ("Parciais", stats["statuses"].get("partial", 0), ORANGE),
        ("Com risco", stats["statuses"].get("risk", 0), RED),
        ("A fazer", stats["statuses"].get("todo", 0), DEEP),
        ("Reservadas", stats["statuses"].get("reserved", 0), MUTED),
    ]
    y = 265
    for title, count, color in rows:
        rounded(draw, (950, y, 1650, y+105), 22, "151E35", color, 2)
        draw.text((990, y+26), title, font=pil_font(24, bold=True), fill=rgb(WHITE))
        draw.text((1530, y+20), str(count), font=pil_font(38, display=True), fill=rgb(color))
        y += 135
    return save_asset(image, "05_progresso.png")


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge, values in edges.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in values.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=110, start=125, bottom=110, end=125):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    prop = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    prop.append(header)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("VOLC O.S.  •  LIVRO VIVO  •  ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText"); text.set(qn("xml:space"), "preserve"); text.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, end))


def style_document(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    specs = [("Title", 34, NAVY), ("Heading 1", 23, NAVY), ("Heading 2", 15, DEEP), ("Heading 3", 11.5, ORANGE)]
    for name, size, color in specs:
        style = doc.styles[name]
        style.font.name = "Space Grotesk"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)


def configure_section(section, cover=False):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    if cover:
        section.top_margin = section.bottom_margin = Cm(0)
        section.left_margin = section.right_margin = Cm(0)
        section.header_distance = section.footer_distance = Cm(0)
        return
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(.7)
    section.footer_distance = Cm(.65)
    header = section.header.paragraphs[0]
    header.text = "VOLC O.S.  •  O LIVRO VIVO DA OPERAÇÃO"
    header.runs[0].font.name = "Space Grotesk"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(DEEP)
    page_number(section.footer.paragraphs[0])


def add_picture(doc: Document, path: Path, width=17.45):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.add_run().add_picture(str(path), width=Cm(width))


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)


def callout(doc: Document, title: str, text: str, color=DEEP, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(.25)
    table.columns[1].width = Cm(16.75)
    set_cell_shading(table.cell(0, 0), color)
    set_cell_shading(table.cell(0, 1), fill)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 150, 165, 150, 165)
        set_cell_border(cell, top={"val":"nil"}, bottom={"val":"nil"}, left={"val":"nil"}, right={"val":"nil"})
    p = table.cell(0, 1).paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True; r.font.name = "Space Grotesk"; r.font.size = Pt(11.5); r.font.color.rgb = RGBColor.from_string(color)
    r = p.add_run(text)
    r.font.size = Pt(9.7); r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def table(doc: Document, headers: list[str], rows: list[tuple], widths: list[float] | None = None, header_color=DEEP, font_size=8.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    if widths:
        for index, width in enumerate(widths):
            t.columns[index].width = Cm(width)
    repeat_header(t.rows[0])
    for index, value in enumerate(headers):
        cell = t.rows[0].cells[index]
        set_cell_shading(cell, header_color)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(value)
        run.bold = True; run.font.size = Pt(7.8); run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, row in enumerate(rows):
        cells = t.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else "F8FAFC")
            set_cell_margins(cell)
            set_cell_border(cell, bottom={"val":"single", "sz":"4", "color":LINE})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def add_inline(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); run.font.name = "Menlo"; run.font.size = Pt(9); run.font.color.rgb = RGBColor.from_string(DEEP)
        else:
            paragraph.add_run(part)


def bullet(doc: Document, text: str, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, text)


def chapter_title(doc: Document, title: str, number: int | None = None):
    if number is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"CAPÍTULO {number:02d}")
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(ORANGE)
    doc.add_heading(title, level=1)


def render_narrative(doc: Document, path: Path, diagrams: dict[str, Path]):
    lines = path.read_text(encoding="utf-8").splitlines()
    h2_number = 0
    paragraph_buffer: list[str] = []
    diagram_for = {
        "A ideia central: operar atenção": ("attention", "Figura 1 — O Mapa da Atenção transforma visão em perguntas operacionais."),
        "As quatro operações que não devem ser misturadas": ("operations", "Figura 2 — Operações diferentes compartilham ativos, não regras."),
        "A fundação comum: Cofre de Ativos": ("vault", "Figura 3 — Famílias iniciais do Cofre de Ativos."),
        "Como ler o roadmap": ("waves", "Figura 4 — Quatro ondas com estacionamento consciente."),
    }

    def flush():
        nonlocal paragraph_buffer
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        if text:
            p = doc.add_paragraph()
            add_inline(p, text)
        paragraph_buffer = []

    first_h2 = True
    for raw in lines:
        line = raw.rstrip()
        if not line:
            flush()
            continue
        if line.startswith("# "):
            continue
        if line.startswith("> "):
            flush()
            callout(doc, "Edição viva", line[2:], PURPLE, PALE_PURPLE)
            continue
        if line.startswith("## "):
            flush()
            if not first_h2:
                doc.add_page_break()
            first_h2 = False
            h2_number += 1
            title = line[3:].strip()
            chapter_title(doc, title, h2_number)
            if title in diagram_for:
                key, text = diagram_for[title]
                add_picture(doc, diagrams[key])
                caption(doc, text)
            continue
        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            flush()
            bullet(doc, line[2:].strip())
            continue
        if re.match(r"^\d+\. ", line):
            flush()
            bullet(doc, re.sub(r"^\d+\. ", "", line), numbered=True)
            continue
        paragraph_buffer.append(line)
    flush()


def capability_next(state: str) -> str:
    return {
        "live": "Preservar, medir e ligar ao restante.",
        "implemented": "Integrar e evitar reconstrução duplicada.",
        "partial": "Fechar o elo e provar a jornada completa.",
        "risk": "Conter, substituir ou governar antes de usar.",
        "todo": "Abrir somente após os pré-requisitos.",
        "historical": "Revalidar antes de reaproveitar.",
        "inactive": "Manter estacionado ou aposentar.",
        "empty": "Decidir se ativa ou aposenta.",
        "decision": "Converter em escolha explícita do dono.",
    }.get(state, "Revisar estado e evidência.")


def add_capability_atlas(doc: Document, graph: dict):
    doc.add_page_break()
    chapter_title(doc, "Atlas das capacidades", 25)
    callout(doc, "Como ler", "Este atlas responde o que cada parte do sistema faz em linguagem simples. Estado não é nota de qualidade: é evidência de maturidade.", CYAN, PALE_BLUE)
    caps = [node for node in graph["nodes"] if node.get("type") == "capability"]
    groups = defaultdict(list)
    for cap in caps:
        groups[cap["cluster"]].append(cap)
    cluster_order = list(graph["clusters"])
    for index, cluster in enumerate(cluster_order):
        if index:
            doc.add_page_break()
        doc.add_heading(graph["clusters"][cluster], level=2)
        rows = []
        for cap in sorted(groups.get(cluster, []), key=lambda item: item["label"]):
            rows.append((
                cap["label"],
                STATE_LABELS.get(cap["state"], cap["state"]),
                cap.get("summary", ""),
                capability_next(cap["state"]),
            ))
        if rows:
            table(doc, ["CAPACIDADE", "ESTADO", "O QUE É", "PRÓXIMO FECHAMENTO"], rows, [3.5, 2.6, 6.3, 4.5], font_size=7.6)


def add_progress_overview(doc: Document, roadmap: dict, stats: dict, diagram: Path):
    doc.add_page_break()
    chapter_title(doc, "Painel geral do Work Road", 35)
    add_picture(doc, diagram)
    caption(doc, "Figura 5 — Índice calculado somente sobre tarefas aceitas; reservadas ficam fora do denominador.")
    callout(doc, "O número certo para esta edição",
            f"{stats['included']} tarefas entram no índice e {stats['reserved']} ficam reservadas. O fechamento editorial atual é {stats['progress']*100:.0f}%.",
            ORANGE, PALE_ORANGE)
    rows = []
    for item in roadmap["initiatives"]:
        progress, included, reserved = progress_for_tasks(item["tasks"], roadmap["status_weights"])
        rows.append((f"{item['rank']:02d}", item["title"], item["wave"], f"{progress*100:.0f}%", f"{included} ativas · {reserved} reservadas"))
    table(doc, ["#", "INICIATIVA", "ONDA", "FECHAMENTO", "ESCOPO"], rows, [1.0, 6.2, 4.0, 2.0, 3.7], font_size=7.5)


def add_initiatives(doc: Document, roadmap: dict):
    for item in roadmap["initiatives"]:
        doc.add_page_break()
        color = WAVE_COLORS.get(item["wave"], DEEP)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"PRIORIDADE {item['rank']:02d}  •  {item['wave'].upper()}")
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(color)
        doc.add_heading(item["title"], level=1)
        progress, included, reserved = progress_for_tasks(item["tasks"], roadmap["status_weights"])
        callout(doc, "Por que existe", item["why"], color, PALE_BLUE if color == CYAN else PALE_ORANGE if color == ORANGE else PALE_PURPLE)
        doc.add_heading("Pronto quando", level=2)
        doc.add_paragraph(item["done_when"])

        rows = []
        for task in item["tasks"]:
            rows.append((
                TASK_ICONS[task["status"]],
                task["title"],
                roadmap["status_labels"][task["status"]],
                task["proof"],
            ))
        table(doc, ["", "TAREFA", "ESTADO", "PROVA OU LACUNA"], rows, [0.8, 6.6, 3.1, 6.4], header_color=color, font_size=7.7)
        doc.add_heading("Leitura do progresso", level=2)
        doc.add_paragraph(f"{progress*100:.0f}% do peso editorial fechado · {included} tarefas no índice · {reserved} reservadas.")
        doc.add_heading("Nós principais do grafo", level=2)
        p = doc.add_paragraph()
        r = p.add_run("  ·  ".join(item["graph_nodes"]))
        r.font.name = "Menlo"; r.font.size = Pt(8.2); r.font.color.rgb = RGBColor.from_string(MUTED)


def add_sources(doc: Document, graph: dict, roadmap: dict):
    doc.add_page_break()
    chapter_title(doc, "Fontes, confiança e atualização", 50)
    callout(doc, "Ordem da verdade",
            "Curadoria humana → snapshot operacional gerado → extração técnica → grafo híbrido → visualizações e livro. O DOCX nunca é fonte.",
            PURPLE, PALE_PURPLE)
    documents = [node for node in graph["nodes"] if node.get("type") == "document"]
    rows = [(node["label"], node.get("cluster_label", ""), node.get("summary", "")) for node in sorted(documents, key=lambda x: x["label"])]
    table(doc, ["DOCUMENTO", "ÁREA", "PAPEL"], rows, [5.2, 4.0, 7.7], font_size=7.4)

    doc.add_page_break()
    doc.add_heading("Checklist de atualização do livro", level=1)
    for text in [
        "Novo material entrou pela inbox estratégica.",
        "Fato, declaração, protótipo, visão e risco foram separados.",
        "A curadoria humana recebeu somente nós e relações aceitos.",
        "O roadmap recebeu tarefa, estado, prova e critério de saída.",
        "O gerador foi executado e o DOCX foi aberto visualmente.",
        "O grafo híbrido foi atualizado pelo pipeline oficial.",
        "Nenhum segredo entrou no livro, JSON, grafo ou log.",
    ]:
        p = doc.add_paragraph()
        p.add_run("☐  ").font.color.rgb = RGBColor.from_string(DEEP)
        p.add_run(text)

    doc.add_heading("Limites desta edição", level=2)
    for text in [
        "O percentual é editorial e não mede prazo, esforço, qualidade ou receita.",
        "A página Facebook é um ativo declarado; sua identidade ainda não foi comprovada.",
        "ChatPion publicado não prova fluxos ou entregabilidade configurados.",
        "Os workflows BEAST estão desativados e nunca operaram de ponta a ponta.",
        "As mudanças de outras frentes continuam preservadas na árvore de trabalho.",
        "Novas estratégias ainda podem ampliar o mapa; a edição é viva por desenho.",
    ]:
        bullet(doc, text)


def add_closing(doc: Document):
    doc.add_page_break()
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(17.3)
    cell = t.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 850, 650, 850, 650)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Você não precisa carregar tudo na cabeça.\n")
    r.font.name = "Space Grotesk"; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(WHITE)
    r = p.add_run("O sistema pode lembrar por você.\n\n")
    r.font.name = "Space Grotesk"; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(CYAN)
    r = p.add_run("VOLC O.S.  •  Livro Vivo da Operação  •  Edição 1.0")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string("B7C3D9")


def build() -> dict:
    graph = json.loads(GRAPH.read_text(encoding="utf-8"))
    roadmap = json.loads(ROADMAP.read_text(encoding="utf-8"))
    stats = roadmap_stats(roadmap)
    diagrams = {
        "cover": draw_cover(stats, graph),
        "attention": draw_attention_map(),
        "operations": draw_operations(),
        "vault": draw_asset_vault(),
        "waves": draw_waves(roadmap),
        "progress": draw_progress(stats),
    }

    doc = Document()
    style_document(doc)
    configure_section(doc.sections[0], cover=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(diagrams["cover"]), width=Cm(21), height=Cm(29.7))

    second = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(second)
    doc.add_heading("Como usar este livro", level=1)
    callout(doc, "Leia em três passadas",
            "Primeiro entenda a operação. Depois veja o atlas e o roadmap. Só então escolha uma prioridade e use seus checklists.",
            ORANGE, PALE_ORANGE)
    table(doc, ["PASSADA", "PERGUNTA", "SAÍDA"], [
        ("1 · Livro", "O que estamos construindo e por quê?", "Clareza sem linguagem técnica."),
        ("2 · Atlas", "O que já existe e em que estado?", "Capacidades e próximos fechamentos."),
        ("3 · Work Road", "O que vem primeiro e como termina?", "Tarefas, prova e percentual editorial."),
    ], [3.1, 7.0, 6.8])
    callout(doc, "A regra mais importante",
            "Itens reservados permanecem no mapa, mas ficam fora do percentual. O futuro não pode fazer o presente parecer fracassado.",
            PURPLE, PALE_PURPLE)

    doc.add_page_break()
    doc.add_heading("Sumário de leitura", level=1)
    headings = [line[3:].strip() for line in NARRATIVE.read_text(encoding="utf-8").splitlines() if line.startswith("## ")]
    for index, title in enumerate(headings, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{index:02d}")
        r.bold = True; r.font.color.rgb = RGBColor.from_string(ORANGE)
        p.add_run(f"   {title}")
    doc.add_paragraph("Parte II · Atlas das capacidades")
    doc.add_paragraph("Parte III · Painel e checklists do Work Road")
    doc.add_paragraph("Parte IV · Fontes, confiança e atualização")

    doc.add_page_break()
    render_narrative(doc, NARRATIVE, diagrams)
    add_capability_atlas(doc, graph)
    add_progress_overview(doc, roadmap, stats, diagrams["progress"])
    add_initiatives(doc, roadmap)
    add_sources(doc, graph, roadmap)
    add_closing(doc)

    doc.core_properties.title = "VOLC O.S. — O Livro Vivo da Operação"
    doc.core_properties.subject = "Visão global, capacidades, prioridades e checklists"
    doc.core_properties.author = "VOLC"
    doc.core_properties.comments = "Gerado a partir do Mapa Vivo e ROADMAP-VIVO.json"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return {
        "output": str(OUT),
        "capabilities": len([node for node in graph["nodes"] if node.get("type") == "capability"]),
        "tasks": stats["total"],
        "reserved": stats["reserved"],
        "editorial_progress": round(stats["progress"] * 100, 1),
        "assets": len(diagrams),
    }


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))
