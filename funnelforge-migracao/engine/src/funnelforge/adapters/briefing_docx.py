"""LEGADO — o caminho ANTIGO do briefing (webgo). Não é mais a origem.

A origem do briefing agora é o `funnel_architecture` do card do Pautador:
ver `adapters/briefing_volc.py` e o comando `funnelforge run-volc`.

Este módulo fica por três motivos, e nenhum deles é sentimento:

1. Ele não é só o leitor de DOCX — é também o leitor de .txt (`read_text`
   abaixo) e é o `loader` de `Deps` montado em `cli.build_deps`. Toda a
   suíte de testes e o comando `run` dependem dele.
2. O comando `run` é o único caminho que já produziu um funil em produção
   (run `antecipacao-saque-aniversario-fgts-20260721-115510`, 7 páginas
   publicadas). Tirá-lo no mesmo commit em que a ponte estreia deixaria a
   ponte sem referência de comparação e sem rollback.
3. O defeito conhecido NÃO é corrigido de propósito: `Document(...)
   .paragraphs` ignora TABELAS, e o briefing do VOLC põe Keywords-alvo,
   CTA, Link do CTA e Slug exatamente numa tabela — foi assim que o funil
   de produção saiu com 0 de 7 páginas com `target_keywords`. Consertar a
   leitura legitimaria um caminho que deve morrer; pela ponte o defeito
   simplesmente não existe, porque nada vira texto.

Quando um card real tiver rodado ponta a ponta por `run-volc`, este arquivo,
o `step_extract` e o `prompts/extractor.jinja` saem juntos — em um commit
separado, que é onde essa remoção pode ser revisada como remoção.
"""
from __future__ import annotations
from pathlib import Path


class DocxBriefingLoader:
    def load(self, path: Path) -> str:
        path = Path(path)
        if path.suffix.lower() == ".docx":
            from docx import Document
            return "\n".join(
                p.text for p in Document(str(path)).paragraphs
            )
        return path.read_text(encoding="utf-8")
